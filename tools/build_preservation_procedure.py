#!/usr/bin/env python3
"""
Generate the Packing, Storage, Preservation & Handling Procedure for SP-00968.

Document: G23BC091-76465-K-KH-0001 / CP-SIF-968-044, Rev A
Project:  Coriolis Fiscal Metering Skid — Statfjord C (Equinor via Apply AS)

Usage:
    python tools/build_preservation_procedure.py
"""

import os
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SHARED_DRIVE = Path(os.environ["USERPROFILE"]) / "OneDrive - Sifab AS" / "Dokumenter - Felles"
PROJECT_DIR = SHARED_DRIVE / "Zigma360" / "Projects" / "SP-00968 Coriolis Skid pakke"
TEMPLATE_PATH = PROJECT_DIR / "05 Dokumentasjon" / "01.Dokumentmaler" / "Front Page.docx"
TARGET_DIR = (
    PROJECT_DIR
    / "05 Dokumentasjon"
    / "02.Dok fra Leverandør"
    / "44. G23BC091-76465-K-KH-0001 (PACKING, STORAGE, PRESERVATION & HANDELING PROCEDURE)"
)
LOGO_PATH = Path(__file__).parent.parent / "temp_logo.jpeg"
OUTPUT_FILENAME = "G23BC091-76465-K-KH-0001_A.docx"

# ---------------------------------------------------------------------------
# Document metadata
# ---------------------------------------------------------------------------
CLIENT_DOC_NO = "G23BC091-76465-K-KH-0001"
SIFAB_DOC_NO = "CP-SIF-968-044"
REVISION = "A"
PO_NUMBER = "4502188671"
PACKAGE_TITLE = "Coriolis Fiscal Metering Skid — Statfjord C"
TAGS = "FT24276, FX24276, PT24276A, PT24276B, TT24276A, TT24276B, TW24276A, TW24276B"
DOC_TITLE = "Packing, Storage, Preservation & Handling Procedure"
TOTAL_PAGES = ""  # Will be set by Word on print


# ---------------------------------------------------------------------------
# Styles & helpers
# ---------------------------------------------------------------------------
SIFAB_GREEN = RGBColor(0x1B, 0x7A, 0x2B)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
TABLE_HEADER_BG = "1B7A2B"
TABLE_ALT_BG = "F2F2F2"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with dict {sz, color, val}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f"<w:tcBorders {nsdecls('w')}/>")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(el)


def add_heading(doc, text, level=1):
    """Add a heading paragraph with manual formatting and outline level for TOC."""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.color.rgb = SIFAB_GREEN
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    # Set outline level so TOC picks it up
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    pPr.append(outline)
    return p


def add_para(doc, text, bold=False, italic=False, size=Pt(10), space_after=Pt(6), alignment=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point with manual bullet character."""
    bullet_char = "\u2022" if level == 0 else "\u2013"
    p = doc.add_paragraph()
    run = p.add_run(f"{bullet_char}  {text}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    return p


def set_table_borders(table):
    """Set borders on all cells in a table (since Table Grid style may not exist)."""
    border_attrs = {"sz": "4", "color": "999999", "val": "single"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                          left=border_attrs, right=border_attrs)


def add_simple_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with green header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for ri_idx, row in enumerate(table.rows):
            for ci_idx, w in enumerate(col_widths):
                row.cells[ci_idx].width = w

    # Apply borders
    set_table_borders(table)

    return table


def add_checklist_table(doc, items, title=None):
    """Add a checklist table with checkbox column."""
    headers = ["#", "Item", "Acceptance Criteria", "Status", "Init./Date"]
    rows = []
    for i, (item, criteria) in enumerate(items, 1):
        rows.append([str(i), item, criteria, "\u2610", ""])
    if title:
        add_para(doc, title, bold=True, size=Pt(11))
    return add_simple_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Build the front page from template
# ---------------------------------------------------------------------------
def build_front_page(doc):
    """Fill in the front page table that was loaded from the template."""
    table = doc.tables[0]

    # Row 4: revision row — first issue
    row4 = table.rows[4]
    row4.cells[0].paragraphs[0].clear()
    run = row4.cells[0].paragraphs[0].add_run(REVISION)
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Issue date
    row4.cells[1].paragraphs[0].clear()
    run = row4.cells[1].paragraphs[0].add_run("2026-03-16")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Description
    for cell in [row4.cells[2], row4.cells[3], row4.cells[4]]:
        cell.paragraphs[0].clear()
    run = row4.cells[2].paragraphs[0].add_run("Issued for Project Acceptance")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Signatures
    row4.cells[5].paragraphs[0].clear()
    run = row4.cells[5].paragraphs[0].add_run("SF")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    row4.cells[7].paragraphs[0].clear()
    run = row4.cells[7].paragraphs[0].add_run("OV")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    row4.cells[8].paragraphs[0].clear()
    run = row4.cells[8].paragraphs[0].add_run("TSF")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Row 6: PO number + Package title
    row6 = table.rows[6]
    for cell in [row6.cells[0], row6.cells[1], row6.cells[2]]:
        cell.paragraphs[0].clear()
    run = row6.cells[0].paragraphs[0].add_run(f"Purchase Order Number:\n{PO_NUMBER}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    for cell in [row6.cells[3]]:
        cell.paragraphs[0].clear()
    run = row6.cells[3].paragraphs[0].add_run(f"Package Title: \n{PACKAGE_TITLE}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Row 8: Manufacturer doc number + rev
    row8 = table.rows[8]
    row8.cells[0].paragraphs[0].clear()
    run = row8.cells[0].paragraphs[0].add_run(f"Manufacturer Document Number:\n {SIFAB_DOC_NO}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    row8.cells[4].paragraphs[0].clear()
    run = row8.cells[4].paragraphs[0].add_run(f"Manufacturer Rev.:\n{REVISION}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Row 9: Manufacturer name/logo
    row9 = table.rows[9]
    row9.cells[0].paragraphs[0].clear()
    p = row9.cells[0].paragraphs[0]
    run = p.add_run("Manufacturer and Supplier Name/ Logo\n")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    if LOGO_PATH.exists():
        run2 = p.add_run()
        run2.add_picture(str(LOGO_PATH), width=Cm(4))

    # Row 10: Tag numbers
    row10 = table.rows[10]
    row10.cells[0].paragraphs[0].clear()
    run = row10.cells[0].paragraphs[0].add_run(f"Tag Number(s):\n{TAGS}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Row 11: Document title
    row11 = table.rows[11]
    row11.cells[0].paragraphs[0].clear()
    run = row11.cells[0].paragraphs[0].add_run(f"Document title:\n{DOC_TITLE}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.bold = True

    # Total pages
    row11.cells[9].paragraphs[0].clear()
    run = row11.cells[9].paragraphs[0].add_run("Total No\nof Pages.")
    run.font.size = Pt(8)
    run.font.name = "Calibri"

    # Row 12: Document number
    row12 = table.rows[12]
    row12.cells[0].paragraphs[0].clear()
    run = row12.cells[0].paragraphs[0].add_run(f"Document Number:\n{CLIENT_DOC_NO}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.bold = True


# ---------------------------------------------------------------------------
# Add header/footer to body pages
# ---------------------------------------------------------------------------
def setup_header_footer(doc):
    """Add header and footer to the document sections (after first page)."""
    # Add a section break for body content
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # Unlink from previous section
    header = section.header
    header.is_linked_to_previous = False
    footer = section.footer
    footer.is_linked_to_previous = False

    # Header: doc number + title
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    run = hp.add_run(f"{CLIENT_DOC_NO}  |  {DOC_TITLE}  |  Rev. {REVISION}")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add bottom border to header
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="808080"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)

    # Footer: company + page number
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    run = fp.add_run("Sifab AS  |  Bedriftsveien 20, 4313 Sandnes  |  ")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # Page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2 = fp.add_run()
    run2._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run3 = fp.add_run()
    run3._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4 = fp.add_run()
    run4._r.append(fldChar2)
    for r in [run2, run3, run4]:
        r.font.size = Pt(8)
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return section


# ---------------------------------------------------------------------------
# SECTION 1: Scope & Purpose
# ---------------------------------------------------------------------------
def section_1(doc):
    add_heading(doc, "1. Scope & Purpose", level=1)

    add_para(doc, (
        "This procedure defines the requirements for packing, storage, preservation, "
        "handling, transportation, and de-preservation of the Coriolis Fiscal Metering Skid "
        "and all associated equipment for the Statfjord C platform (Equinor), supplied by "
        "Sifab AS under Apply AS Purchase Order."
    ))

    add_para(doc, (
        "The procedure applies from completion of Factory Acceptance Testing (FAT) at "
        "Sifab AS workshop in Sandnes, through packing, marine transport, offshore crane "
        "lifting, platform storage, and de-preservation for final installation and commissioning."
    ))

    add_heading(doc, "1.1 Equipment Covered", level=2)

    equipment_table = [
        ["FX24276", "Micro Motion ELITE Coriolis Flow Sensor", "1", "SS 316L"],
        ["FT24276", "Micro Motion 5700 Transmitter", "1", "Aluminium / SS 316"],
        ["PT24276A", "Pressure Transmitter", "1", "SS 316L"],
        ["PT24276B", "Pressure Transmitter", "1", "SS 316L"],
        ["TT24276A", "Temperature Transmitter", "1", "SS 316L"],
        ["TT24276B", "Temperature Transmitter", "1", "SS 316L"],
        ["TW24276A", "Thermowell", "1", "SS 316L"],
        ["TW24276B", "Thermowell", "1", "SS 316L"],
        ["—", "DBB Valve Assembly (with pneumatic actuators)", "As per P&ID", "SS 316 / Duplex"],
        ["—", "GRP Instrument Enclosure", "1", "GRP (fibreglass)"],
        ["—", "Piping, Pipe Spools & Flanged Connections", "As per isometric", "SS 316L / Duplex"],
        ["—", "Skid Frame (structural steel)", "1", "Carbon steel (coated per M-501)"],
        ["—", "Electrical Cables, Cable Trays & Junction Boxes", "As per cable schedule", "Various"],
    ]

    add_simple_table(doc,
        ["Tag No.", "Description", "Qty", "Material"],
        equipment_table,
    )

    add_para(doc, "")  # spacer

    add_heading(doc, "1.2 Applicability", level=2)
    add_para(doc, "This procedure shall be applied at the following lifecycle stages:")
    stages = [
        "Workshop preservation — after FAT completion and before packing",
        "Packing and crating — securing equipment for transport",
        "Road transport — from Sifab AS workshop to supply base",
        "Marine transport — supply vessel to Statfjord C platform",
        "Offshore crane lift — from supply vessel deck to platform receiving area",
        "Platform storage — temporary storage pending installation",
        "De-preservation — removal of protective measures prior to installation and commissioning",
    ]
    for s in stages:
        add_bullet(doc, s)


# ---------------------------------------------------------------------------
# SECTION 2: References
# ---------------------------------------------------------------------------
def section_2(doc):
    add_heading(doc, "2. References", level=1)

    refs = [
        ["Norsok Z-001", "Documentation for Operation", "Documentation, marking, and traceability"],
        ["Norsok Z-015", "Temporary Equipment", "Temporary protective measures during installation"],
        ["Norsok M-501", "Surface Preparation and Protective Coating", "Coating systems, inspection, repair (§5–§12)"],
        ["Norsok I-001", "Field Instrumentation", "Instrument storage, preservation, installation (§10)"],
        ["Norsok R-002", "Lifting Equipment", "Lifting lug design, proof load testing, SWL marking"],
        ["Norsok R-003", "Safe Use of Lifting Equipment", "Lift planning, rigging, sling angles, weather limits"],
        ["Norsok E-001", "Electrical Installations", "Electrical preservation, halogen-free materials, megger testing"],
        ["Norsok S-001", "Technical Safety", "General safety requirements"],
        ["ISPM-15", "Int'l Standards for Phytosanitary Measures No. 15", "Heat-treated wood for international shipping"],
        ["ISO 780", "Packaging — Pictorial Marking for Handling of Goods", "Handling symbols for crates"],
        ["ISO 8501-1", "Surface Preparation Grades", "Visual standards Sa 2.5, Sa 3"],
        ["Equinor TR2000", "Equinor Technical Requirements", "Equinor-specific packaging and preservation"],
        ["Equinor TR0052", "Equinor Requirements for Preservation", "Offshore preservation requirements"],
        ["Emerson 20002158", "Installation Manual — Micro Motion ELITE Sensors", "Handling, lifting, and mounting"],
        ["Emerson 20002166", "Installation Manual — Micro Motion 5700 Transmitters", "Wiring, grounding, and protection"],
        ["B.SI.01.07", "Sifab AS — Procedure for Document Handling", "QMS document control"],
    ]

    add_simple_table(doc,
        ["Document", "Title", "Relevance"],
        refs,
    )


# ---------------------------------------------------------------------------
# SECTION 3: Definitions & Abbreviations
# ---------------------------------------------------------------------------
def section_3(doc):
    add_heading(doc, "3. Definitions & Abbreviations", level=1)

    defs = [
        ["CoG", "Centre of Gravity"],
        ["DBB", "Double Block and Bleed"],
        ["DFT", "Dry Film Thickness"],
        ["FAT", "Factory Acceptance Test"],
        ["GRP", "Glass Reinforced Plastic (fibreglass)"],
        ["ISPM-15", "International Standards for Phytosanitary Measures No. 15"],
        ["MRB", "Manufacturing Record Book"],
        ["PE", "Polyethylene"],
        ["PM", "Project Manager"],
        ["PO", "Purchase Order"],
        ["QA/QC", "Quality Assurance / Quality Control"],
        ["SMDL", "Supplier Master Document List"],
        ["SS 316L", "Austenitic Stainless Steel Grade 316L (UNS S31603)"],
        ["SWL", "Safe Working Load"],
        ["TSA", "Thermal Spray Aluminium"],
        ["VCI", "Volatile Corrosion Inhibitor"],
    ]

    add_simple_table(doc,
        ["Abbreviation", "Definition"],
        defs,
    )


# ---------------------------------------------------------------------------
# SECTION 4: Responsibilities
# ---------------------------------------------------------------------------
def section_4(doc):
    add_heading(doc, "4. Responsibilities", level=1)

    roles = [
        ["Sifab AS Project Manager", (
            "Overall responsibility for execution of this procedure. Ensures all preservation, "
            "packing, and handling activities are performed per this document. Reviews and approves "
            "preservation logs before shipment."
        )],
        ["Sifab AS Workshop / Fabrication", (
            "Performs all workshop preservation, packing, crating, and loading activities. "
            "Records all preservation activities on the checklists in Section 12. Ensures "
            "carbon steel separation per Norsok M-501 §5 during final assembly."
        )],
        ["Sifab AS QA/QC", (
            "Verifies that preservation materials, coating, and packing comply with this procedure "
            "and referenced standards. Witnesses critical hold points. Signs off preservation logs."
        )],
        ["Apply AS (Contractor)", (
            "Reviews and approves this procedure. Coordinates offshore receiving, crane operations, "
            "and platform storage arrangements. Provides platform storage location requirements."
        )],
        ["Equinor (Client)", (
            "Reviews and approves this procedure per SMDL requirements. Provides platform-specific "
            "storage and handling constraints."
        )],
        ["Platform Lifting Coordinator", (
            "Responsible for offshore crane lift operations per Norsok R-003. Verifies lift plan, "
            "weather conditions, and rigging arrangements. Issues lifting permits."
        )],
    ]

    for role, desc in roles:
        add_heading(doc, role, level=3)
        add_para(doc, desc)


# ---------------------------------------------------------------------------
# SECTION 5: General Preservation Requirements
# ---------------------------------------------------------------------------
def section_5(doc):
    add_heading(doc, "5. General Preservation Requirements", level=1)

    add_heading(doc, "5.1 Storage Conditions", level=2)
    add_para(doc, "All equipment shall be stored under the following conditions unless otherwise specified:")
    conditions = [
        "Indoor, dry, well-ventilated storage preferred. If outdoor storage is unavoidable, full weatherproof covering shall be provided.",
        "Temperature: –10\u00b0C to +50\u00b0C (storage). Equipment containing electronics shall not be stored below –20\u00b0C.",
        "Relative humidity: < 60% RH preferred. Humidity indicators shall be placed inside all sealed packages.",
        "Equipment shall be stored clear of the ground on pallets or dunnage, minimum 100 mm above floor level.",
        "Storage area shall be free from vibration, chemical fumes, salt spray, and direct sunlight.",
    ]
    for c in conditions:
        add_bullet(doc, c)

    add_heading(doc, "5.2 Material Compatibility — Critical Requirements", level=2)
    add_para(doc, (
        "The following material compatibility requirements are mandatory and are based on "
        "lessons learned from SP-00525/577 (NOA / Valhall Fenris prover projects), where "
        "incorrect material grades resulted in extensive corrosion remediation (USD 3,900+ "
        "per unit and 33+ labour hours)."
    ), bold=False, italic=True)

    add_para(doc, "Chloride Control on Stainless Steel:", bold=True)
    cl_items = [
        "No chloride-containing materials (PVC tape, chlorinated solvents, HCl-based flux) shall contact stainless steel surfaces.",
        "All marking pens, adhesive tapes, and protective films used on SS surfaces shall be certified chloride-free.",
        "Cleaning agents shall be chloride-free and suitable for use on austenitic stainless steel.",
    ]
    for item in cl_items:
        add_bullet(doc, item)

    add_para(doc, "Material Grade Requirements (Lessons Learned — SP-00525/577):", bold=True)
    mat_items = [
        "ALL stainless steel fasteners, cable ties, seal wires, washers, and ancillary components shall be SS 316 (NOT SS 304). This applies to every BOM line item without exception.",
        "Material certificates (EN 10204 Type 3.1) shall be verified at FAT for all stainless steel components.",
        "Carbon steel components (rivets, bolts, washers) are NOT acceptable for offshore North Sea service unless specifically approved and coated per Norsok M-501.",
        "Edge protection strips, seal wires, and similar small components are particularly susceptible — verify material grade before installation.",
    ]
    for item in mat_items:
        add_bullet(doc, item)

    add_heading(doc, "5.3 Approved Preservation Materials", level=2)

    pres_materials = [
        ["Acid-free vaseline (petroleum jelly)", "Machined surfaces, threads, flange faces", "Norsok I-001 §10"],
        ["VCI (Volatile Corrosion Inhibitor) paper/bags", "Loose fasteners, spare parts, enclosed spaces", "MIL-PRF-22019"],
        ["Silica gel desiccant sachets", "Inside sealed packages and crates", "MIL-D-3464"],
        ["Humidity indicator cards (HIC)", "Inside all sealed packages", "MIL-I-8835"],
        ["PE (polyethylene) shrink wrap / vapour barrier", "Overall equipment wrapping", "—"],
        ["ISPM-15 heat-treated wood", "Crate construction, pallets, dunnage", "ISPM-15"],
        ["Chloride-free protective film", "Stainless steel plate/pipe surfaces", "—"],
        ["Norsok M-501 approved topcoat (touch-up)", "Coating repairs during transport/storage", "Norsok M-501 §11"],
    ]

    add_simple_table(doc,
        ["Material", "Application", "Specification"],
        pres_materials,
    )

    add_heading(doc, "5.4 Workshop Contamination Control", level=2)
    add_para(doc, (
        "Based on SP-00525/577 lessons learned, where carbon steel workshop debris "
        "contaminated painted motor and speed reducer surfaces:"
    ))
    ws_items = [
        "Carbon steel fabrication and grinding shall not be performed in the same area as skid final assembly (Norsok M-501 §5).",
        "Final clean-down shall be performed with stainless-safe solvents before any preservation is applied.",
        "Workshop floor and workbenches shall be cleaned before placing the skid for preservation activities.",
        "Personnel shall use clean gloves when handling preserved components to prevent fingerprint contamination.",
    ]
    for item in ws_items:
        add_bullet(doc, item)


# ---------------------------------------------------------------------------
# SECTION 6: Equipment-Specific Preservation
# ---------------------------------------------------------------------------
def section_6(doc):
    add_heading(doc, "6. Equipment-Specific Preservation", level=1)

    add_para(doc, (
        "This section defines the specific preservation requirements for each equipment "
        "category on the metering skid. All preservation shall be completed after FAT and "
        "before packing (Section 7)."
    ))

    # 6.1 Coriolis meters
    add_heading(doc, "6.1 Coriolis Meters (Micro Motion ELITE + 5700 Transmitters)", level=2)
    add_para(doc, "Tags: FX24276 (sensor), FT24276 (transmitter)", bold=True)

    add_para(doc, "Sensor (ELITE):", bold=True)
    items = [
        "Install factory-supplied flange protectors on all process connections. If original protectors are unavailable, use steel blind flanges with neoprene gaskets, secured with minimum 4 bolts per flange (ref. Emerson 20002158, Chapter 3).",
        "Apply acid-free vaseline to all exposed flange sealing faces.",
        "Protect sensor case vent/purge connections with plastic caps. Do NOT seal purge ports airtight — allow pressure equalisation per Emerson guidelines.",
        "Wrap sensor flow tubes in PE protective film, secured with chloride-free tape.",
        "Do NOT lift the sensor by the junction box, transmitter housing, or purge fittings (ref. Emerson 20002158, §3.1).",
        "The sensor may be lifted by its case using continuous web belt slings or eye-to-eye slings (ref. Emerson 20002158, Figure 3-1).",
        "Place VCI sachets inside the PE wrapping near the sensor body.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_para(doc, "Transmitter (5700):", bold=True)
    items = [
        "Install protective caps on all conduit entries and cable glands.",
        "Verify all conduit entries face downward to prevent moisture ingress. If orientation cannot be maintained during transport, seal entries with waterproof tape over the caps.",
        "Wrap transmitter housing in PE film with desiccant sachets inside.",
        "Protect display window with adhesive protective film.",
        "Disconnect and cap any field wiring cables. Seal cable ends with heat-shrink tubing.",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.2 DBB Valves
    add_heading(doc, "6.2 DBB Valves with Pneumatic Actuators", level=2)
    items = [
        "Close all valves to the fully closed position.",
        "Apply acid-free vaseline to all exposed stem threads and packing areas.",
        "Install blind flanges or steel end caps with neoprene gaskets on all open process connections. Secure with minimum 4 bolts each.",
        "Protect pneumatic actuator air supply connections with plastic caps.",
        "Wrap pneumatic actuator cylinders in PE film to prevent moisture and salt spray ingress.",
        "Apply VCI paper wrapping to all exposed carbon steel components of actuator brackets.",
        "Tag each valve with its tag number on a stainless steel wire tag (SS 316).",
        "Ensure all valve position indicators are visible through preservation wrapping or clearly marked on external labelling.",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.3 Pressure transmitters
    add_heading(doc, "6.3 Pressure Transmitters & Manifolds", level=2)
    add_para(doc, "Tags: PT24276A, PT24276B", bold=True)
    items = [
        "Install factory-supplied process connection caps on all pressure ports.",
        "Apply acid-free vaseline to all exposed threaded connections.",
        "Seal conduit entries with protective caps and waterproof tape.",
        "Wrap each transmitter in PE film with a desiccant sachet.",
        "For 3-valve or 5-valve manifolds: close all valves, apply vaseline to all stem threads, cap all open ports.",
        "If transmitters are mounted on the skid: ensure mounting brackets do not impose mechanical stress during transport. Add temporary bracing if required.",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.4 Temperature sensors & thermowells
    add_heading(doc, "6.4 Temperature Sensors & Thermowells", level=2)
    add_para(doc, "Tags: TT24276A, TT24276B, TW24276A, TW24276B", bold=True)
    items = [
        "Apply acid-free vaseline to thermowell process threads and flange sealing faces.",
        "Install protective caps on thermowell bores to prevent debris ingress.",
        "Wrap temperature sensor heads in PE film. Seal conduit entries with caps.",
        "If sensors are pre-installed in thermowells: apply additional wrapping around the sensor/thermowell junction.",
        "Place a desiccant sachet in each wrapped assembly.",
        "Verify thermowell material certificates confirm SS 316L (per lessons learned — SS 304 is NOT acceptable).",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.5 GRP enclosure
    add_heading(doc, "6.5 GRP Instrument Enclosure", level=2)
    items = [
        "Close and latch all enclosure doors/panels.",
        "Seal all cable entry points with temporary foam plugs or plastic caps.",
        "Place desiccant sachets (minimum 4 per enclosure) inside the enclosure with humidity indicator cards.",
        "Seal door/panel edges with weatherproof tape to prevent moisture ingress during transport.",
        "Protect external GRP surfaces with PE film to prevent UV degradation during outdoor storage.",
        "Do not stack anything on top of the GRP enclosure.",
        "Verify all internal terminal strips, DIN rails, and cable terminations are secure and labelled.",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.6 Piping
    add_heading(doc, "6.6 Piping, Pipe Spools & Flanged Connections", level=2)
    items = [
        "All open pipe ends shall be sealed with steel blind flanges (with neoprene gaskets) or welded steel caps. Plastic caps are acceptable only for temporary workshop protection — NOT for shipment.",
        "Apply acid-free vaseline to all flange sealing faces (raised face, ring-type joint grooves).",
        "Coat all flange bolting (studs, nuts, washers) with acid-free vaseline. All bolting shall be SS 316 grade.",
        "For ring-type joint (RTJ) flanges: install the service gasket in the groove, protected with vaseline, and secure the blind flange over it.",
        "Internal pipe surfaces shall be clean and dry. Blow through with dry, oil-free compressed air or nitrogen before sealing.",
        "Apply VCI sachets inside each sealed pipe section (insert through a drain valve or vent connection before final sealing).",
        "All weld areas shall be visually inspected for coating damage before preservation. Touch up per Norsok M-501 §11.",
        "Remove all temporary welding attachments. Grind smooth and touch up coating.",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.7 Skid frame
    add_heading(doc, "6.7 Skid Frame (Structural Steel)", level=2)
    items = [
        "The skid frame coating system shall be per Norsok M-501 System 1 for carbon steel in atmospheric Zone 1 exposure. All coated surfaces shall be verified per the project Coating Report (G23BC091-76465-K-RV-0003).",
        "Inspect all coated surfaces for damage (scratches, chips, abrasion) after FAT. Record all damage on the coating inspection log.",
        "Touch up all coating damage per Norsok M-501 §11 requirements. Surface preparation for touch-up shall be to minimum St 3 (power tool cleaning) or Sa 2.5 (blast cleaning) depending on damage extent.",
        "Allow full coating cure time before packing or wrapping (minimum 7 days for epoxy systems, or as specified by coating manufacturer).",
        "Apply additional protective wrapping to areas likely to suffer transport damage: skid corners, lifting lug areas, and forklift contact points.",
        "All exposed unpainted surfaces (lifting lugs, forklift pockets, jacking points) shall be coated with acid-free vaseline.",
        "Verify SWL markings on lifting lugs remain clearly visible after all preservation activities (Norsok R-002 §10).",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6.8 Electrical
    add_heading(doc, "6.8 Electrical Cables, Cable Trays & Junction Boxes", level=2)
    items = [
        "All cable ends shall be sealed with heat-shrink tubing or self-amalgamating tape. Cable glands shall be fitted with dust caps.",
        "Junction boxes: close and secure all covers. Seal cable entry points with protective caps. Place desiccant sachets inside each box.",
        "Verify all cable trays are secure and will not shift during transport. Apply temporary cable ties (SS 316 only) at 300 mm intervals where needed.",
        "All cables shall be halogen-free per Norsok E-001 requirements.",
        "Wrap exposed cable runs in PE film for protection against mechanical damage during transport.",
        "Protect all earthing/bonding connection points from corrosion with acid-free vaseline. Verify continuity before sealing.",
        "Record insulation resistance (megger test) values for all cables before preservation. These baseline values shall be compared with de-preservation readings (Section 11).",
        "Cable labels and identification tags shall remain visible or be duplicated on the external preservation wrapping.",
    ]
    for item in items:
        add_bullet(doc, item)


# ---------------------------------------------------------------------------
# SECTION 7: Packing & Crating
# ---------------------------------------------------------------------------
def section_7(doc):
    add_heading(doc, "7. Packing & Crating", level=1)

    add_heading(doc, "7.1 Crate Construction", level=2)
    items = [
        "All wooden crate materials shall be ISPM-15 compliant heat-treated wood, marked with the IPPC/HT stamp on all four sides. This is mandatory for international shipments (ref. QCP 006 §2.3).",
        "Crate design shall be sufficient for the combined weight of the skid plus all packaged equipment. Design factor of safety: minimum 3:1 for static loads.",
        "Crate base (skid) shall be bolted (not nailed) to the skid frame at minimum 4 points using M16 or larger bolts with washers.",
        "Crate side panels and top cover shall be secured with lag bolts (minimum 3\" × 3/8\" lag bolts with washers).",
        "Cross-banding shall be applied on all six sides of the crate for structural rigidity.",
        "Crate shall include forklift pockets (if applicable) or provisions for crane lifting with slings.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "7.2 Vapour Barrier & Desiccant", level=2)
    items = [
        "After all equipment-specific preservation (Section 6) is complete, wrap the entire skid in PE stretch wrap liner.",
        "Place an encapsulating Mylar or aluminium-laminated vapour barrier bag over the wrapped skid.",
        "Heat-seal all bag seams to the bottom liner, leaving one opening for vacuum extraction.",
        "Vacuum extract air from the bag to maximum extent practicable. Heat-seal the final opening.",
        "Distribute silica gel desiccant sachets throughout the crate floor and around the equipment. Minimum quantity: 1 kg desiccant per cubic metre of enclosed volume.",
        "Place humidity indicator cards (minimum 2 per crate) at visible positions (viewable through an inspection window or at crate opening points).",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "7.3 Shock & Tilt Indicators", level=2)
    items = [
        "Install minimum two (2) shock indicators (ShockWatch or equivalent) on opposite crate surfaces.",
        "Install minimum two (2) tilt indicators (TiltWatch or equivalent) on adjacent crate surfaces.",
        "Record indicator serial numbers and initial status on the packing checklist (Section 12).",
        "Photograph installed indicators before shipment as baseline evidence.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "7.4 Securing for Marine Transport", level=2)
    items = [
        "The crated skid shall be secured to the transport vehicle / supply vessel deck using ratchet straps or chain binders rated for the skid weight plus a dynamic load factor of 2.0.",
        "Securing arrangements shall prevent movement in all directions: longitudinal, transverse, and vertical.",
        "D-rings or lashing eyes on the crate base shall be provided for tie-down attachment.",
        "For supply vessel transport: the crate shall be positioned on deck per the vessel's cargo plan and secured in accordance with the vessel master's requirements and IMO CSS Code.",
        "Anti-skid material (rubber matting or similar) shall be placed between the crate base and the deck surface.",
    ]
    for item in items:
        add_bullet(doc, item)


# ---------------------------------------------------------------------------
# SECTION 8: Marking & Labeling
# ---------------------------------------------------------------------------
def section_8(doc):
    add_heading(doc, "8. Marking & Labelling", level=1)

    add_para(doc, (
        "All marking shall comply with Norsok Z-001 and ISO 780 requirements. "
        "Markings shall be applied on all four vertical sides of the crate using "
        "weather-resistant paint or stencil."
    ))

    add_heading(doc, "8.1 Required Markings", level=2)

    markings = [
        ["Consignee", "Equinor — Statfjord C"],
        ["Destination", "Statfjord C Platform, North Sea"],
        ["Supplier", "Sifab AS, Bedriftsveien 20, 4313 Sandnes, Norway"],
        ["PO Number", PO_NUMBER],
        ["Tag Numbers", TAGS],
        ["Package/Crate ID", "SP-00968-CRATE-001"],
        ["Document Reference", CLIENT_DOC_NO],
        ["Gross Weight", "[To be completed] kg"],
        ["Net Weight", "[To be completed] kg"],
        ["Dimensions (L × W × H)", "[To be completed] mm × mm × mm"],
        ["Centre of Gravity (CoG)", "Indicated by \u2295 symbol on two adjacent sides"],
        ["Lifting Points", "Indicated by \u21c5 symbol with SWL"],
    ]

    add_simple_table(doc, ["Item", "Content"], markings)

    add_heading(doc, "8.2 Handling Symbols (per ISO 780)", level=2)
    symbols = [
        "\"This Way Up\" arrows on all four sides",
        "\"Fragile — Handle With Care\" symbol",
        "\"Keep Dry\" (umbrella) symbol",
        "\"Sling Here\" indicators adjacent to approved lifting points",
        "\"Do Not Stack\" symbol on top panel",
        "\"Temperature Limits\" if applicable",
    ]
    for s in symbols:
        add_bullet(doc, s)


# ---------------------------------------------------------------------------
# SECTION 9: Handling & Lifting
# ---------------------------------------------------------------------------
def section_9(doc):
    add_heading(doc, "9. Handling & Lifting", level=1)

    add_heading(doc, "9.1 General Requirements", level=2)
    items = [
        "All lifting and handling operations shall comply with Norsok R-002 (lifting equipment design) and Norsok R-003 (safe use of lifting equipment).",
        "A lift plan shall be prepared for each lift operation (workshop loading, supply base handling, offshore crane lift).",
        "All lifting equipment (slings, shackles, spreader beams) shall be certified and colour-coded per the current inspection period.",
        "The lift plan shall include: equipment weight, CoG location, rigging arrangement, sling angles, SWL verification, and environmental conditions.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "9.2 Workshop Handling", level=2)
    items = [
        "Use overhead crane or forklift for workshop handling. Forklift forks shall be rubber-padded or wrapped to prevent coating damage.",
        "The skid shall be lifted using designated lifting lugs only. Lifting lugs shall be proof-load tested and certified per Norsok R-002 before first use.",
        "Do not drag or slide the skid on any surface.",
        "During loading onto transport vehicle: use certified slings through lifting lugs. Minimum sling angle: 60° from horizontal.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "9.3 Road Transport", level=2)
    items = [
        "Transport vehicle shall be suitable for the crate dimensions and weight.",
        "Secure the crate per Section 7.4.",
        "Maximum transport speed on public roads: as per local regulations.",
        "Driver shall be briefed on the fragile nature of the cargo and instructed to avoid sudden acceleration, braking, and sharp turns.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "9.4 Supply Vessel Transport", level=2)
    items = [
        "The crate shall be positioned per the vessel cargo plan.",
        "Securing shall comply with IMO CSS Code and vessel master's requirements.",
        "Weather covers (tarpaulin or equivalent) shall be provided if the crate is not fully weatherproof.",
        "The crate shall not be placed where it may be subjected to wave impact, cargo shifting, or crane operations during transit.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "9.5 Offshore Crane Lift", level=2)
    add_para(doc, (
        "The offshore crane lift is the highest-risk handling operation. All requirements "
        "of Norsok R-003 shall apply."
    ))
    items = [
        "A detailed lift plan shall be prepared and approved by the Platform Lifting Coordinator before the lift.",
        "Lift plan shall include: gross weight (skid + crate + packaging), CoG, rigging arrangement, crane capacity chart verification, exclusion zone, and communication plan.",
        "Weather limits: maximum wind speed for lifting shall be per platform operating procedures (typically 12 m/s sustained wind). No lifting in conditions exceeding the crane's rated capacity at the required radius.",
        "Sling angles: minimum 60° from horizontal. Spreader beam to be used if sling angles cannot be maintained.",
        "Tag lines shall be used to control the load during the lift. Minimum two tag lines, operated by trained riggers.",
        "Personnel exclusion zone shall be established below and around the lift path.",
        "Radio communication between crane operator, banksman, and deck crew shall be maintained throughout the lift.",
        "The load shall not be suspended over personnel or occupied areas at any time.",
    ]
    for item in items:
        add_bullet(doc, item)


# ---------------------------------------------------------------------------
# SECTION 10: Storage on Platform
# ---------------------------------------------------------------------------
def section_10(doc):
    add_heading(doc, "10. Storage on Platform", level=1)

    add_heading(doc, "10.1 Storage Location Requirements", level=2)
    items = [
        "The skid shall be stored in a designated laydown area approved by the platform operations team.",
        "Storage location shall be on a level, load-rated surface capable of supporting the full crate weight.",
        "The crate shall be positioned to allow access for inspection and re-preservation activities.",
        "Avoid storage locations exposed to direct wave spray, exhaust gases, or chemical discharge.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "10.2 Inspection Intervals", level=2)

    intervals = [
        ["First 30 days", "Weekly", "Visual inspection of crate condition, shock/tilt indicators, humidity indicators (if visible)"],
        ["30 days – 6 months", "Bi-weekly", "As above, plus check for condensation, water pooling, crate structural integrity"],
        ["6 – 12 months", "Monthly", "Full inspection: open crate, check humidity indicators, replace desiccants if >60% RH, inspect coating, check VCI effectiveness"],
        ["> 12 months", "Monthly", "Full re-preservation may be required. Contact Sifab AS for re-preservation procedure."],
    ]

    add_simple_table(doc,
        ["Storage Duration", "Inspection Frequency", "Scope"],
        intervals,
    )

    add_heading(doc, "10.3 Maximum Storage Duration", level=2)
    add_para(doc, (
        "The preservation as specified in this procedure is designed for a maximum "
        "storage period of 12 months from the date of packing. If installation is "
        "delayed beyond 12 months, Sifab AS shall be contacted for assessment and "
        "potential re-preservation. Extended storage may require:"
    ))
    items = [
        "Complete re-inspection of all preservation measures",
        "Replacement of all desiccants and VCI materials",
        "Coating inspection and touch-up per Norsok M-501 §11",
        "Re-measurement of insulation resistance for all cables",
        "Updated preservation log entries",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "10.4 Humidity Monitoring", level=2)
    add_para(doc, (
        "Humidity indicator cards shall be checked at each inspection. If any indicator "
        "shows > 60% RH, the crate shall be opened and desiccants replaced immediately. "
        "The cause of moisture ingress shall be investigated (damaged vapour barrier, "
        "broken seal, etc.) and corrected before re-sealing."
    ))


# ---------------------------------------------------------------------------
# SECTION 11: De-Preservation Procedure
# ---------------------------------------------------------------------------
def section_11(doc):
    add_heading(doc, "11. De-Preservation Procedure", level=1)

    add_para(doc, (
        "De-preservation shall be performed immediately prior to installation. "
        "Do not de-preserve equipment that will not be installed within 48 hours."
    ))

    add_heading(doc, "11.1 Uncrating", level=2)
    items = [
        "Inspect crate exterior for damage. Record status of shock and tilt indicators on the de-preservation checklist (Section 12).",
        "Photograph the crate and all indicators before opening.",
        "Remove crate top cover and side panels carefully. Do not use pry bars directly against the equipment.",
        "Inspect the vapour barrier for damage. Record any tears, punctures, or moisture ingress evidence.",
        "Check humidity indicator cards. Record readings on the de-preservation checklist.",
        "Carefully remove the vapour barrier and PE wrapping. Collect and dispose of all desiccant sachets.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "11.2 Instrument De-Preservation", level=2)
    add_para(doc, "Applies to: Coriolis meters, pressure transmitters, temperature sensors, thermowells.", bold=True)
    items = [
        "Remove all PE wrapping, VCI materials, and protective films.",
        "Remove protective caps from all process connections, conduit entries, and cable glands.",
        "Remove acid-free vaseline from all flange faces and threaded connections using a clean, lint-free cloth with approved solvent (e.g., isopropyl alcohol).",
        "Inspect all sealing surfaces for damage, corrosion, or contamination. Record any findings.",
        "Verify tag numbers are correct and legible on all instruments.",
        "For Coriolis sensor (FX24276): remove flange protectors. Inspect flow tube bore for debris or contamination. Clean if required per Emerson guidelines.",
        "For transmitter (FT24276): remove conduit caps. Verify display is functional (power-up test if feasible).",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "11.3 Valve De-Preservation", level=2)
    items = [
        "Remove all PE wrapping, blind flanges, and VCI materials from DBB valves and actuators.",
        "Remove vaseline from all stem threads and packing areas.",
        "Operate each valve through full stroke (open-close-open) to verify free movement.",
        "For pneumatic actuators: remove caps from air supply connections. Connect temporary air supply and verify actuator operation.",
        "Inspect all valve body and actuator surfaces for corrosion or coating damage.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "11.4 Piping De-Preservation", level=2)
    items = [
        "Remove all blind flanges and steel end caps.",
        "Remove vaseline from all flange sealing faces using clean, lint-free cloths with approved solvent.",
        "Inspect internal pipe surfaces (visually and with a strong light) for corrosion, debris, or contamination.",
        "Blow through all pipe sections with dry, oil-free compressed air or nitrogen.",
        "Inspect all weld areas for coating damage. Touch up per Norsok M-501 §11 if required.",
        "Install service gaskets on all flanged connections per the project specification.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "11.5 Electrical De-Preservation", level=2)
    items = [
        "Remove all PE wrapping, caps, and sealing tape from cable ends, junction boxes, and cable glands.",
        "Remove desiccants from junction boxes. Inspect for moisture or corrosion. Record findings.",
        "Perform insulation resistance (megger) test on all cables. Compare results with baseline values recorded before preservation (Section 6.8). Minimum acceptable insulation resistance: 1 M\u03a9 at 500 VDC.",
        "Perform continuity test on all earthing/bonding connections. Verify < 1 \u03a9 resistance.",
        "Inspect all cable labels and identification tags. Replace any damaged or illegible labels.",
        "Remove vaseline from all earthing connection points.",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "11.6 Coating Inspection", level=2)
    items = [
        "Perform visual inspection of all coated surfaces per Norsok M-501.",
        "Measure DFT at representative locations. Compare with the project Coating Report values.",
        "Record all coating damage on the coating inspection log.",
        "Touch up all damage areas per Norsok M-501 §11. Allow specified cure time before installation.",
        "For areas requiring full re-coating: contact Sifab AS for coating repair procedure.",
    ]
    for item in items:
        add_bullet(doc, item)


# ---------------------------------------------------------------------------
# SECTION 12: Preservation Log & Checklists
# ---------------------------------------------------------------------------
def section_12(doc):
    add_heading(doc, "12. Preservation Log & Checklists", level=1)

    add_para(doc, (
        "The following checklists shall be completed during preservation, packing, and "
        "de-preservation activities. Completed checklists shall be included in the "
        "Manufacturing Record Book (MRB)."
    ))

    # Checklist A: Preservation Execution
    add_heading(doc, "12.1 Checklist A — Preservation Execution", level=2)
    pres_items = [
        ("Coriolis sensor (FX24276): flange protectors installed, vaseline applied, PE wrapped, VCI placed", "Visual inspection"),
        ("Coriolis transmitter (FT24276): conduit caps installed, PE wrapped, desiccant placed", "Visual inspection"),
        ("Pressure transmitters (PT24276A/B): process caps installed, vaseline applied, PE wrapped", "Visual inspection"),
        ("Temperature sensors (TT24276A/B): PE wrapped, conduit entries capped, desiccant placed", "Visual inspection"),
        ("Thermowells (TW24276A/B): vaseline on threads/faces, bore caps installed, material cert verified SS 316L", "Material cert check"),
        ("DBB valves: closed position, vaseline on stems, blind flanges on process connections, actuators wrapped", "Visual inspection"),
        ("GRP enclosure: doors latched, cable entries sealed, desiccants placed (min. 4), humidity indicators installed", "Visual + count"),
        ("All piping: steel blind flanges installed, vaseline on flange faces, internal surfaces dry, VCI inside", "Visual inspection"),
        ("All bolting: vaseline applied, material verified SS 316", "Material cert check"),
        ("Skid frame: coating inspected post-FAT, all damage touched up per M-501 §11, lifting lugs preserved", "DFT measurement"),
        ("Cables: heat-shrink on all ends, junction boxes sealed with desiccants, megger baseline recorded", "Megger test record"),
        ("Cable ties: all verified SS 316 (NOT SS 304)", "Material cert check"),
        ("Workshop contamination control: CS separation confirmed, final clean-down completed", "Visual inspection"),
        ("All stainless components: material grade verified SS 316 (not SS 304) per SP-00525/577 lessons", "Material cert check"),
        ("All preservation materials: chloride-free, approved per Section 5.3", "Data sheet check"),
        ("Vapour barrier installed and heat-sealed", "Visual inspection"),
        ("Desiccant quantity: minimum 1 kg per m\u00b3 enclosed volume", "Calculation + count"),
        ("Humidity indicator cards installed (min. 2 per crate)", "Visual + count"),
        ("Shock indicators installed (min. 2) — serial numbers recorded", "Serial number log"),
        ("Tilt indicators installed (min. 2) — serial numbers recorded", "Serial number log"),
    ]
    add_checklist_table(doc, pres_items)
    add_para(doc, "")

    # Sign-off block
    signoff = [
        ["Prepared by:", "", "Date:", ""],
        ["Checked by:", "", "Date:", ""],
        ["Approved by:", "", "Date:", ""],
    ]
    add_simple_table(doc, ["Role", "Name / Signature", "", ""], signoff)

    # Checklist B: Periodic Inspection Log
    add_heading(doc, "12.2 Checklist B — Periodic Storage Inspection Log", level=2)
    add_para(doc, "One row per inspection visit. Attach as continuation sheet if more rows are needed.")

    insp_headers = ["Date", "Inspector", "Crate Condition", "Shock Ind.", "Tilt Ind.", "Humidity %", "Action Taken", "Sign."]
    insp_rows = [["", "", "", "", "", "", "", ""] for _ in range(12)]
    add_simple_table(doc, insp_headers, insp_rows)

    # Checklist C: De-Preservation
    add_heading(doc, "12.3 Checklist C — De-Preservation Completion", level=2)
    depres_items = [
        ("Crate exterior inspected — shock/tilt indicators status recorded and photographed", "Photo evidence"),
        ("Humidity indicators checked — readings recorded", "Reading < 60% = OK"),
        ("Vapour barrier and PE wrapping removed and inspected for damage", "Visual inspection"),
        ("All instruments: preservation removed, sealing surfaces inspected, tag numbers verified", "Visual inspection"),
        ("Coriolis sensor bore inspected for debris/contamination", "Visual inspection"),
        ("All valves: operated through full stroke, actuators function-tested", "Functional test"),
        ("All piping: blind flanges removed, internal surfaces inspected, blown through with dry air/N\u2082", "Visual + air test"),
        ("All flange faces: vaseline removed, sealing surfaces inspected, service gaskets installed", "Visual inspection"),
        ("Coating: visual inspection completed, DFT spot-checked, damage areas touched up", "DFT measurement"),
        ("Cables: megger test completed, results compared with baseline, all > 1 M\u03a9", "Megger test record"),
        ("Earthing/bonding: continuity tested, all connections < 1 \u03a9", "Continuity test record"),
        ("All cable labels verified legible and correct", "Visual inspection"),
        ("All preservation waste collected and disposed of properly", "Visual confirmation"),
        ("De-preservation completion form signed by all parties", "Signatures below"),
    ]
    add_checklist_table(doc, depres_items)
    add_para(doc, "")

    # Sign-off block
    add_simple_table(doc, ["Role", "Name / Signature", "", ""], signoff)


# ---------------------------------------------------------------------------
# SECTION 13: Appendices
# ---------------------------------------------------------------------------
def section_13(doc):
    add_heading(doc, "13. Appendices", level=1)

    add_heading(doc, "Appendix A — Equipment Tag List", level=2)
    tags_table = [
        ["FX24276", "Micro Motion ELITE Coriolis Flow Sensor", "SS 316L", "Sifab AS"],
        ["FT24276", "Micro Motion 5700 Transmitter", "Aluminium / SS 316", "Sifab AS"],
        ["PT24276A", "Pressure Transmitter", "SS 316L", "Sifab AS"],
        ["PT24276B", "Pressure Transmitter", "SS 316L", "Sifab AS"],
        ["TT24276A", "Temperature Transmitter", "SS 316L", "Sifab AS"],
        ["TT24276B", "Temperature Transmitter", "SS 316L", "Sifab AS"],
        ["TW24276A", "Thermowell", "SS 316L", "Sifab AS"],
        ["TW24276B", "Thermowell", "SS 316L", "Sifab AS"],
    ]
    add_simple_table(doc, ["Tag No.", "Description", "Material", "Supplier"], tags_table)

    add_heading(doc, "Appendix B — Lifting Sketch Reference", level=2)
    add_para(doc, (
        "Refer to the project General Arrangement drawing for lifting lug locations, "
        "Safe Working Load (SWL) values, and Centre of Gravity (CoG) position. "
        "The following information shall be confirmed from the GA drawing before any lift:"
    ))
    items = [
        "Number and location of lifting lugs",
        "SWL per lifting lug (kg or tonnes)",
        "Skid gross weight including all mounted equipment",
        "CoG location (longitudinal, transverse, and vertical offsets from a datum point)",
        "Minimum sling length and maximum sling angle from horizontal",
        "Recommended rigging arrangement (direct sling, spreader beam, etc.)",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "Appendix C — Approved Preservation Materials List", level=2)
    materials = [
        ["Acid-free vaseline", "e.g., Shell Ensis, Fuchs Anticorit", "Machined surfaces, threads, flanges"],
        ["VCI paper (vapour phase)", "e.g., Cortec VpCI-146, Zerust ICT", "Wrapping ferrous/non-ferrous metals"],
        ["VCI emitter sachets", "e.g., Cortec VpCI-101, Zerust ActivPak", "Inside sealed packages"],
        ["Silica gel desiccant", "e.g., Clariant Desi Pak, SorbentSystems", "Moisture absorption in crates"],
        ["Humidity indicator cards", "Per MIL-I-8835 (6-spot)", "Monitoring enclosed humidity"],
        ["PE stretch wrap", "200 \u03bcm min. thickness, UV-stabilised", "Overall equipment wrapping"],
        ["Vapour barrier film (Mylar/Al)", "e.g., Cortec MilCorr, Royco barrier", "Encapsulating wrapped equipment"],
        ["Chloride-free protective tape", "e.g., 3M 471, Advance AT7", "Securing wrapping on SS surfaces"],
        ["Coating touch-up paint", "Per project Coating Report", "Norsok M-501 §11 repair"],
    ]
    add_simple_table(doc,
        ["Material", "Brand / Specification", "Application"],
        materials,
    )

    add_heading(doc, "Appendix D — Crate Marking Diagram", level=2)
    add_para(doc, "The crate shall be marked on all four vertical sides as shown below:")
    add_para(doc, "")

    # ASCII-style crate marking diagram
    diagram_text = (
        "    ┌─────────────────────────────────────────┐\n"
        "    │                                         │\n"
        "    │   CONSIGNEE: Equinor — Statfjord C      │\n"
        "    │   SUPPLIER:  Sifab AS, Sandnes, Norway  │\n"
        "    │   PO:        " + PO_NUMBER + "                │\n"
        "    │   TAGS:      FX24276, FT24276, etc.     │\n"
        "    │   CRATE ID:  SP-00968-CRATE-001         │\n"
        "    │   GROSS WT:  ______ kg                  │\n"
        "    │   DIMS:      ____ × ____ × ____ mm      │\n"
        "    │                                         │\n"
        "    │   [↑ THIS WAY UP]  [☂ KEEP DRY]        │\n"
        "    │   [⚠ FRAGILE]      [⊕ CoG HERE]        │\n"
        "    │   [⊘ DO NOT STACK]                      │\n"
        "    │                                         │\n"
        "    └─────────────────────────────────────────┘"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Table of Contents placeholder
# ---------------------------------------------------------------------------
def add_toc(doc):
    """Add a Table of Contents field that Word will populate on open."""
    add_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run("[Right-click and select 'Update Field' to populate Table of Contents]")
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run4.font.size = Pt(10)
    run4.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("Loading front page template...")
    doc = Document(str(TEMPLATE_PATH))

    print("Filling in front page...")
    build_front_page(doc)

    # Add section break and set up body page header/footer
    print("Setting up body pages...")
    setup_header_footer(doc)

    # Table of Contents
    add_toc(doc)

    # Document sections
    print("Writing Section 1: Scope & Purpose...")
    section_1(doc)

    print("Writing Section 2: References...")
    section_2(doc)

    print("Writing Section 3: Definitions & Abbreviations...")
    section_3(doc)

    print("Writing Section 4: Responsibilities...")
    section_4(doc)

    print("Writing Section 5: General Preservation Requirements...")
    section_5(doc)

    print("Writing Section 6: Equipment-Specific Preservation...")
    section_6(doc)

    print("Writing Section 7: Packing & Crating...")
    section_7(doc)

    print("Writing Section 8: Marking & Labelling...")
    section_8(doc)

    print("Writing Section 9: Handling & Lifting...")
    section_9(doc)

    print("Writing Section 10: Storage on Platform...")
    section_10(doc)

    print("Writing Section 11: De-Preservation Procedure...")
    section_11(doc)

    print("Writing Section 12: Preservation Log & Checklists...")
    section_12(doc)

    print("Writing Section 13: Appendices...")
    section_13(doc)

    # Save
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TARGET_DIR / OUTPUT_FILENAME
    doc.save(str(output_path))
    print(f"\nDocument saved to:\n  {output_path}")
    print(f"\nTotal sections: 13")
    print(f"Document: {CLIENT_DOC_NO} / {SIFAB_DOC_NO}, Rev. {REVISION}")

    return output_path


if __name__ == "__main__":
    path = main()
