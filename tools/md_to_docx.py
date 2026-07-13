"""Markdown -> DOCX converter for Sifab project documents.

Targets the structure used in deviation lists, OCs and similar:
- # / ## / ### headings
- GitHub-flavoured tables (| ... |)
- Bullet lists (- ...)
- Numbered lists (1. ...)
- Blockquotes (> ...)
- Horizontal rules (---)
- Inline **bold** and *italic*

Usage:
    python tools/md_to_docx.py <input.md> <output.docx>
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SIFAB_GREEN = RGBColor(0x2A, 0x5D, 0x2A)


def add_runs_with_inline_formatting(paragraph, text, *, base_size=11, color=None):
    """Add runs to a paragraph honouring **bold**, *italic* and `code`."""
    if not text:
        return
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(base_size)
        run.font.name = run.font.name or "Calibri"
        if color is not None:
            run.font.color.rgb = color


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tc_pr.append(shd)


def add_horizontal_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0C0C0")
    pBdr.append(bottom)
    p_pr.append(pBdr)


def parse_table_rows(lines):
    """Parse |a|b|c| style lines into list of cell-lists."""
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        # Strip leading/trailing pipes and split
        inner = line.strip()
        if inner.startswith("|"):
            inner = inner[1:]
        if inner.endswith("|"):
            inner = inner[:-1]
        cells = [c.strip() for c in inner.split("|")]
        rows.append(cells)
    return rows


def is_table_separator(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def render(md_text, output_path):
    doc = Document()

    # Set page size to A4 landscape if needed; default A4 portrait is fine for deviation lists
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Tweak default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Configure heading colours / sizes
    for level, size in [(1, 22), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = SIFAB_GREEN

    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1; continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            add_horizontal_rule(p)
            i += 1; continue

        # Empty line
        if not stripped:
            i += 1; continue

        # Tables
        if stripped.startswith("|"):
            table_block = []
            while i < n and lines[i].strip().startswith("|"):
                table_block.append(lines[i])
                i += 1
            # Filter separator
            data_lines = [ln for ln in table_block if not is_table_separator(ln)]
            rows = parse_table_rows(data_lines)
            if rows:
                n_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=n_cols)
                tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(n_cols):
                        cell = tbl.cell(ri, ci)
                        cell.text = ""
                        text = row[ci] if ci < len(row) else ""
                        para = cell.paragraphs[0]
                        add_runs_with_inline_formatting(para, text, base_size=10)
                        if ri == 0:
                            shade_cell(cell, "2A5D2A")
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.bold = True
                        elif ri % 2 == 0:
                            shade_cell(cell, "EFF6EF")
                doc.add_paragraph()  # spacing
            continue

        # Blockquote
        if stripped.startswith("> "):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                ln = lines[i].strip()[1:].strip()
                quote_lines.append(ln)
                i += 1
            quote_text = " ".join(quote_lines).strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_runs_with_inline_formatting(p, quote_text, base_size=11)
            continue

        # Bullet list
        if stripped.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                content = lines[i].strip()[2:]
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_inline_formatting(p, content, base_size=11)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                content = re.sub(r"^\d+\.\s", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_runs_with_inline_formatting(p, content, base_size=11)
                i += 1
            continue

        # Plain paragraph
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].startswith("#") or lines[i].strip().startswith("|")
            or lines[i].strip().startswith(">") or lines[i].strip().startswith("- ")
            or re.match(r"^\d+\.\s", lines[i].strip()) or lines[i].strip() == "---"
        ):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs_with_inline_formatting(p, " ".join(para_lines), base_size=11)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python tools/md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    render(src.read_text(encoding="utf-8"), str(dst))
    print(f"Wrote {dst}  ({dst.stat().st_size/1024:.1f} KB)")
