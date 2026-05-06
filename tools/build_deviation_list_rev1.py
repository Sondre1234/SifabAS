"""Generate SP-01415 Deviation List Rev 1 (post-PO 4500998501) as Word document.

Re-issued at PO acceptance per PO clause: "any deviation to this Purchase Order
(technical/commercial), this must be made in a separate document and issued to
Guidant upon receipt of this order."

Reuses the 18 deviations from Rev 0 with status updates and adds DEV-019 for
the Pos 4 (dis/re-assembly hourly work) VOR mechanism.
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

import sys
sys.path.insert(0, str(Path(__file__).parent))
from build_deviation_list import DEVIATIONS as REV0_DEVIATIONS, set_cell, add_header_table, severity_color  # noqa: E402

from _paths import SHARED
DEST = SHARED / 'Zigma360' / 'Projects' / 'SP-01415 Small Volume Prover Snorre A' / '05 Dokumentasjon' / '04.Dok sendt til kunde'

# --------------------------------------------------------------------------
# Status updates per deviation (Rev 0 → Rev 1)
# --------------------------------------------------------------------------
STATUS_UPDATES = {
    'DEV-001': 'Open — Honeywell to confirm IP66 alternatives.',
    'DEV-002': 'In progress — Honeywell (S. Swart) verbally confirmed 60 Hz at no cost (23 March 2026). Awaiting written confirmation.',
    'DEV-003': 'Open — formal deviation request to be submitted (Aker BP precedent).',
    'DEV-004': 'In progress — Sifab to source NORSOK M-501 System 6C painting locally in Norway. Cost included in quote.',
    'DEV-005': 'Open — Honeywell EWA quote required to align with 28-month warranty from platform installation.',
    'DEV-006': 'Open — Honeywell to confirm high-temperature option (insulation plate + jacket) is included.',
    'DEV-007': 'In progress — verbally accepted by Guidant on 18 March 2026. Formal deviation request with split-module drawings to follow within 4 WAO.',
    'DEV-008': 'Open — Honeywell supervisor rates confirmed (USD 140 onshore / USD 180 offshore + KPI). SAT scope to be quoted.',
    'DEV-009': 'Open — Honeywell to confirm 316/316L compliance with TR2000 BD20X MDS. 6Mo tubing/fittings will be Sifab free-issue.',
    'DEV-010': 'Open — Honeywell to clarify PMI scope per TR1427.',
    'DEV-011': 'In progress — Sifab to supply NORSOK-compliant cables, glands and cable trays as free-issue (Aker BP precedent).',
    'DEV-012': 'Open — gap analysis between Honeywell standard package and RFQ §1.16 to be completed.',
    'DEV-013': 'Open — to be included in Sifab’s back-to-back PO to Honeywell.',
    'DEV-014': 'Open — Sifab engineering scope. NORSOK R-002 lifting calculations to be developed.',
    'DEV-015': 'Open — engineering review during GA drawing phase.',
    'DEV-016': 'Open — Honeywell to confirm hydrotest pressure meets 1.5 × BD20X maximum design pressure (≈ 73.5 barg).',
    'DEV-017': 'In progress — Sifab to coordinate Justervesenet certification with Pemberton lead time.',
    'DEV-018': 'Open — Sifab to align Honeywell payment milestones with Guidant Milestone 1 (clarified PO + bank guarantee). No Sifab pre-financing.',
}

# --------------------------------------------------------------------------
# New DEV-019 — Pos 4 hourly work / VOR mechanism
# --------------------------------------------------------------------------
NEW_DEVIATIONS = [
    {
        'id': 'DEV-019',
        'title': 'Pos 4 Onshore Disassembly + Offshore Reinstallation – Hourly Work and VOR Mechanism',
        'severity': 'MEDIUM',
        'rfq_ref': 'Sifab Quote SP-01415, Pos 4 / PO 4500998501',
        'requirement': 'Onshore disassembly, packing, transportation, and offshore reinstallation of the SVP at Snorre A, including Honeywell Enraf supervisor and Intertek prover expert support, performed on an hourly basis.',
        'offered': 'In Sifab’s quotation, Pos 4 was included as a budget estimate (approximately USD 110,000) on hourly basis only, with rates of USD 140/hour onshore and USD 180/hour offshore (plus KPI adjustment) for both Honeywell Enraf supervisor and Intertek prover expert. Number of mobilisations/demobilisations to be advised based on availability and waiting time.',
        'deviation': 'Purchase Order 4500998501 consolidates the entire scope into a single line item (Item 00010, USD 1,608,262.36). The hourly-based Pos 4 work is included in scope but is not separately itemised. The actual cost depends on field conditions, weather, mobilisation logistics and platform access at Snorre A.',
        'risk_owner': 'Sifab + Guidant (cost variation)',
        'resolution': 'Sifab proposes that the actual hourly cost for Pos 4 (onshore disassembly + offshore reinstallation, supervised by Honeywell Enraf and Intertek) is reconciled and invoiced under the Variation Order (VOR) mechanism per PP-PS-13, against pre-agreed rates: USD 140/hour onshore and USD 180/hour offshore (+ KPI adjustment), plus mobilisation/demobilisation as advised. A baseline forecast of approximately USD 110,000 is estimated; any variation against the budget will be transparently documented and submitted as a VOR for Guidant approval.',
    },
]


def build_rev1():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    doc.add_heading('Deviation List', level=0)
    doc.add_heading('SP-01415 Snorre A – Small Volume Prover', level=1)

    # Header
    add_header_table(doc, [
        ('RFQ No.', 'GM-8501-1447'),
        ('Sifab Project', 'SP-01415'),
        ('Customer Project', 'GM-5341 Snorre A Oil Metering'),
        ('Purchase Order', '4500998501 (issued 29 April 2026)'),
        ('Customer', 'Guidant – Measurement Solutions Norway AS'),
        ('End Client', 'Equinor – Snorre A platform'),
        ('Honeywell Proposal', '10465986-O-1010834 R0, dated 20 March 2026'),
        ('Prepared by', 'Sifab AS'),
        ('Date', '5 May 2026'),
        ('Revision', '1'),
    ])

    # Opening note
    doc.add_paragraph('')
    note = doc.add_paragraph()
    r = note.add_run('Note: ')
    r.bold = True
    note.add_run(
        'This Deviation List Rev 1 is issued in response to Purchase Order 4500998501 in '
        'accordance with the PO clause requiring deviations to be documented in a separate '
        'document and submitted to Guidant upon receipt of this order. Sifab AS acts as system '
        'integrator and intermediary between Guidant and Honeywell Enraf. Deviations are '
        'documented to ensure full transparency. Sifab does not accept technical risk for items '
        'outside its scope — risk sits with the party responsible for the deliverable. Items '
        'previously documented in Rev 0 (23 March 2026) are carried forward with updated '
        'status. One additional commercial deviation (DEV-019) has been added covering the '
        'hourly basis of Pos 4 (onshore disassembly + offshore reinstallation work).'
    )

    # All deviations: Rev 0 + new
    all_devs = list(REV0_DEVIATIONS) + NEW_DEVIATIONS

    for dev in all_devs:
        doc.add_page_break()
        doc.add_heading(f"{dev['id']}: {dev['title']}", level=2)

        sev_p = doc.add_paragraph()
        sev_run = sev_p.add_run(f"Severity: {dev['severity']}")
        sev_run.bold = True
        sev_run.font.color.rgb = severity_color(dev['severity'])

        status = STATUS_UPDATES.get(dev['id'], 'Open')

        t = doc.add_table(rows=7, cols=2)
        t.style = 'Light Grid Accent 1'
        fields = [
            ('RFQ / PO Reference', dev['rfq_ref']),
            ('Requirement', dev['requirement']),
            ('Honeywell Offer / Quote Position', dev['offered']),
            ('Deviation', dev['deviation']),
            ('Risk Owner', dev['risk_owner']),
            ('Proposed Resolution', dev['resolution']),
            ('Status (Rev 1)', status),
        ]
        for i, (k, v) in enumerate(fields):
            set_cell(t.rows[i].cells[0], k, bold=True)
            set_cell(t.rows[i].cells[1], v)

    # Summary
    doc.add_page_break()
    doc.add_heading('Summary', level=1)

    high = sum(1 for d in all_devs if d['severity'] == 'HIGH')
    med = sum(1 for d in all_devs if d['severity'] == 'MEDIUM')
    low = sum(1 for d in all_devs if d['severity'] == 'LOW')

    p = doc.add_paragraph()
    p.add_run(
        f'HIGH severity items: {high}  |  MEDIUM: {med}  |  LOW: {low}  |  Total: {len(all_devs)}'
    )

    st = doc.add_table(rows=len(all_devs) + 1, cols=4)
    st.style = 'Light Grid Accent 1'
    for i, h in enumerate(['#', 'Deviation', 'Severity', 'Status (Rev 1)']):
        set_cell(st.rows[0].cells[i], h, bold=True)
    for i, dev in enumerate(all_devs):
        set_cell(st.rows[i + 1].cells[0], dev['id'])
        set_cell(st.rows[i + 1].cells[1], dev['title'])
        set_cell(st.rows[i + 1].cells[2], dev['severity'])
        # short status
        long_status = STATUS_UPDATES.get(dev['id'], 'Open')
        short = 'Open' if long_status.lower().startswith('open') else 'In progress'
        set_cell(st.rows[i + 1].cells[3], short)

    # Save
    DEST.mkdir(parents=True, exist_ok=True)
    outpath = DEST / 'SP-01415_Deviation_List_Rev1.docx'
    doc.save(str(outpath))
    print(f'Saved: {outpath}')


if __name__ == '__main__':
    build_rev1()
