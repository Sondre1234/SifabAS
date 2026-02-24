"""
Fill QCF524 Specification Worksheet for Snorre A SVP085.
Rule: NEVER remove existing text. Only ADD values alongside.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from copy import deepcopy

# Start from the ORIGINAL template
src = 'projects/snorre-a-compact-prover/honeywell/QCF524_Specification Worksheet Rev V19.docx'
dst = 'projects/snorre-a-compact-prover/engineering/QCF524_Snorre_A_FILLED.docx'

doc = Document(src)


def add_to_cell(table_idx, row_idx, col_idx, value):
    """Add value to a cell WITHOUT removing existing text."""
    cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
    existing = cell.text.strip()
    if existing:
        # Append to existing text
        cell.paragraphs[0].text = existing + "  " + value
    else:
        # Empty cell — just add value
        cell.paragraphs[0].text = value


def set_cell(table_idx, row_idx, col_idx, value):
    """Set cell to a specific value (for cells that are meant to be filled in)."""
    cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
    cell.paragraphs[0].text = value


# ============================================================
# Table 0: Meter type — mark Ultrasonic with X
# ============================================================
add_to_cell(0, 4, 1, "X")                          # Ultrasonic -> check
add_to_cell(0, 4, 2, "Krohne Altosonic 5, 8\" CL300")  # Maker
add_to_cell(0, 4, 3, "67 m\u00b3/h")               # Min flow
add_to_cell(0, 4, 4, "750 m\u00b3/h")              # Max flow
add_to_cell(0, 4, 5, "~400 m\u00b3/h")             # Normal flow

# ============================================================
# Table 1: Operating and Design temperature
# Original text: "Range: Min      Max      Normal"
# ADD values after the template text
# ============================================================
add_to_cell(1, 0, 1, "  55\u00b0C / 57\u00b0C / 56\u00b0C")
add_to_cell(1, 0, 3, "X")  # Mark degrees C column

add_to_cell(1, 1, 1, "  -8\u00b0C / 106\u00b0C")
add_to_cell(1, 1, 3, "X")  # Mark degrees C column

# ============================================================
# Table 2: Operating and Design pressure
# Original: "Range: Min      Max      Normal\nRange: Min      Max"
# ============================================================
add_to_cell(2, 0, 1, "  35.4 / 36.0 / 35.7 barg (oper.) / -8 / 49 barg (design)")
add_to_cell(2, 0, 4, "X")  # Mark bar column

# ============================================================
# Table 3: Viscosity, Density, Flow rate
# Original has labels only — fill in the values
# ============================================================
add_to_cell(3, 0, 0, "  2.2 cP")
add_to_cell(3, 0, 1, "  815 kg/m\u00b3")
add_to_cell(3, 0, 2, "  67-750 m\u00b3/h")

# ============================================================
# Table 5: Type of installation — Offshore, 316L
# Row 3 = Offshore prover
# ============================================================
add_to_cell(5, 3, 2, "X")  # 316/316L column

# ============================================================
# Table 6: Flange Rating — Class 600 RJ (row 7, code 8C)
# The "Y" in the SVP085 column (col 7) already exists
# Add X next to it in the SVP085 column
# ============================================================
add_to_cell(6, 7, 7, " <-- SELECTED")  # Mark 8C / Class 600 RJ / SVP085

# ============================================================
# Table 7: Flange orientation — "06" = both on top
# Row 4: keep existing text, add X in check column
# ============================================================
add_to_cell(7, 4, 2, "X")  # Check column (col 2)

# ============================================================
# Table 9: ATEX voltage — 220/240 VAC 50Hz 3Ph (row 3)
# But RFQ says 230VAC 60Hz — add note
# Mark row 3 for SVP085 (col 7) as closest match + note
# ============================================================
add_to_cell(9, 3, 7, "X (TBC - RFQ states 230VAC 3ph 60Hz, confirm with Guidant)")

# ============================================================
# Table 10: Hazardous area — ATEX II 2(1)G
# Row 2: ATEX classification
# ============================================================
add_to_cell(10, 2, 1, "  X (Zone 1, IIA T3 min per RFQ - HW std IIB T4 exceeds)")

# ============================================================
# Table 11: Paint — Other (row 2)
# ============================================================
add_to_cell(11, 2, 1, "X - NORSOK M-501 System 6C per TR0042")

# ============================================================
# Table 12: Calibration — Gravimetric (row 0)
# ============================================================
add_to_cell(12, 0, 1, "X")

# ============================================================
# Table 14: Insulation — Jacket and plate (row 2)
# ============================================================
add_to_cell(14, 2, 1, "X - Prepared for min 100mm insulation")

# ============================================================
# Table 15: NACE — Yes (row 1)
# ============================================================
add_to_cell(15, 1, 1, "X")

# ============================================================
# Table 16: Seal — Carbon Fiber PTFE (row 1)
# ============================================================
add_to_cell(16, 1, 1, "X")

# ============================================================
# Table 18: Transmitters — No transmitters (row 3)
# ============================================================
add_to_cell(18, 3, 1, "X (TE free-issued by Guidant, PT by Guidant)")

# ============================================================
# Table 19: PED (row 1)
# ============================================================
add_to_cell(19, 1, 1, "X")

# ============================================================
# Table 20: PMI (row 1)
# ============================================================
add_to_cell(20, 1, 1, "X - 10% SS316, 100% Duplex/6Mo per TR1427")

# ============================================================
# Table 22: Packing — Other / Modular split (row 3)
# ============================================================
add_to_cell(22, 3, 1, "X - Modular split, max W1.4 x L2.56 x H2.2m per module")

# ============================================================
# Table 23: Optional items
# ============================================================
add_to_cell(23, 4, 1, "X (option - 2yr spare parts)")     # Spare parts
add_to_cell(23, 5, 1, "X (re-assembly + commissioning on Snorre A)")  # Commissioning
add_to_cell(23, 7, 1, "X (Norsok compliance required)")    # DNV offshore

# ============================================================
# Table 25: Flow computer — Krohne
# ============================================================
add_to_cell(25, 0, 3, " (X) - Altosonic 5")  # Krohne column

# ============================================================
# Table 28: Contacts — fill in empty fields
# ============================================================
set_cell(28, 0, 1, "Sondre Falch")
set_cell(28, 0, 3, "+47 900 29 588")
set_cell(28, 0, 7, "sondre.falch@sifab.no")

set_cell(28, 1, 1, "Tom Sverre Falch")
set_cell(28, 1, 3, "+47 416 28 408")
set_cell(28, 1, 7, "tom.falch@sifab.no")

set_cell(28, 2, 1, "Tom Sverre Falch")
set_cell(28, 2, 3, "+47 416 28 408")
set_cell(28, 2, 7, "tom.falch@sifab.no")

set_cell(28, 3, 1, "Equinor ASA")
set_cell(28, 3, 3, "Norway - Snorre A Platform")


doc.save(dst)
print(f"Saved: {dst}")
print("Done - all original text preserved, values added alongside.")
