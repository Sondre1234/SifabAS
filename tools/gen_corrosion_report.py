#!/usr/bin/env python3
"""Generate Corrosion Analysis Report for SP-00525/SP-00577."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import datetime

BLUE = RGBColor(0, 51, 102)
RED = RGBColor(180, 0, 0)
GREY = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)

SHARED = os.path.join(os.environ["USERPROFILE"], "OneDrive - Sifab AS", "Dokumenter - Felles")
LOGO_PATH = os.path.join(SHARED, "Logoer", "SIFAB", "GB SIFAB LOGO GRØNN 03.jpg")
PROJECT_DIR = os.path.join(SHARED, "Zigma360", "Projects",
                           "SP-00525 and SP-00577 Prover NOA og Valhall Fenris")

today = datetime.date.today().strftime("%d.%m.%Y")


def shade_row(row, fill_color):
    for cell in row.cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill_color
        })
        tc_pr.append(shd)


def set_cell(cell, text, bold=False, size=9, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_heading_colored(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE
    return h


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── TITLE PAGE ──
    if os.path.exists(LOGO_PATH):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(LOGO_PATH, width=Cm(5))
        except Exception:
            pass

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CORROSION ANALYSIS REPORT")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE

    for text, size in [
        ("SP-00525 (Hugin A / NOA) & SP-00577 (Valhall Fenris)", 14),
        ("Honeywell Enraf Small Volume Provers — SVP050 & SVP015", 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = GREY

    doc.add_paragraph()

    # Info table
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate([
        ("Document No.", "SIF-525/577-RPT-001"),
        ("Date", today),
        ("Revision", "Rev 0"),
        ("Prepared by", "Sondre Falch, Sales Manager"),
        ("Client", "Aker BP ASA (via TechnipFMC)"),
        ("Classification", "Internal — For Warranty Claim / Lessons Learned"),
    ]):
        set_cell(info_table.rows[i].cells[0], k, bold=True, size=10)
        set_cell(info_table.rows[i].cells[1], v, size=10)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──
    add_heading_colored(doc, "1. Executive Summary")
    doc.add_paragraph(
        "This report documents the corrosion and material non-conformance issues identified on two "
        "Honeywell Enraf Small Volume Provers delivered as part of projects SP-00525 (Hugin A / NOA) "
        "and SP-00577 (Valhall Fenris PWP). The provers were manufactured by TruStop Inc. (Phoenix, AZ) "
        "and delivered to Sifab AS for integration, testing, and final delivery to Aker BP via TechnipFMC."
    )
    doc.add_paragraph(
        "During pre-delivery inspection, end user review (Aker BP), and Mecan inspection (20 November 2024), "
        "significant corrosion was found on multiple components across both units. The root cause is a "
        "material specification gap \u2014 the fabricator used SS 304 and carbon steel fasteners where SS 316 "
        "was required for North Sea offshore service. This was compounded by insufficient surface "
        "protection, workshop contamination, and incomplete painting."
    )
    doc.add_paragraph(
        "All identified corrosion issues were remediated prior to deployment. Corrective actions included "
        "replacement of affected parts with SS 316 components and application of acid-free vaseline "
        "preservation. Total remediation cost: approximately USD 12,000 (108.5 hours + parts)."
    )

    # ── 2. SCOPE ──
    add_heading_colored(doc, "2. Scope & Equipment")

    t2 = doc.add_table(rows=7, cols=3)
    t2.style = "Light Grid Accent 1"
    for j, h in enumerate(["Parameter", "Valhall PWP (SVP015)", "Hugin A (SVP050)"]):
        set_cell(t2.rows[0].cells[j], h, bold=True, size=9)
    for i, row_data in enumerate([
        ["Project No.", "SP-00577", "SP-00525"],
        ["Tag", "17-IH-19121", "D-21MQ001"],
        ["Prover Model", 'PR15-0031 (6")', 'PR50-0070 (8")'],
        ["PO No.", "4500747127", "4500747089"],
        ["Service", "Fenris condensate metering", "NOA fiscal oil export"],
        ["Location", "Valhall platform", "Hugin A platform"],
    ]):
        for j, val in enumerate(row_data):
            set_cell(t2.rows[i + 1].cells[j], val, size=9)

    # ── 3. CORROSION FINDINGS ──
    doc.add_paragraph()
    add_heading_colored(doc, "3. Corrosion Findings \u2014 Part-by-Part Analysis")
    doc.add_paragraph(
        "The following table lists all components where corrosion was identified, "
        "the original material, the failure mode, and the corrective action taken."
    )

    ft = doc.add_table(rows=1, cols=6)
    ft.style = "Table Grid"
    ft_headers = ["#", "Component", "Unit(s)", "Original Material", "Failure Mode", "Corrective Action"]
    for j, h in enumerate(ft_headers):
        set_cell(ft.rows[0].cells[j], h, bold=True, size=8, color=WHITE)
    shade_row(ft.rows[0], "003366")

    findings = [
        ["1", "Pop rivets\n(motor compartment cover)", "Valhall", "Carbon steel / plated",
         "Corroded through plating; visible rust on all rivets",
         "Full panel replacement; all rivets upgraded to SS 316"],
        ["2", "Pillow blocks", "Valhall", "Painted carbon steel",
         "Corrosion on unpainted surface areas",
         "Pillow blocks replaced; grease applied to unpainted surfaces"],
        ["3", "Shock absorbers", "Both", "Plated carbon steel",
         "Corrosion on body and threads",
         "Replaced with new units; acid-free vaseline preservation"],
        ["4", "End flange tool bolt", "Both", "Coated carbon steel",
         "Bolt rusted up \u2014 tool unusable",
         "Tool removed (space limitation); custom replacement tool planned"],
        ["5", "End flange coated bolts", "Both", "Coated carbon steel",
         "Coating failed; corrosion underneath",
         "Bolts replaced; acid-free vaseline applied"],
        ["6", "Optical switch seal wire", "Both", "Non-316 stainless wire",
         "Wire susceptible to chloride corrosion",
         "Changed to SS 316 wire"],
        ["7", "Motor surface", "Both", "Painted (factory finish)",
         "Rust debris contamination from workshop",
         "Cleaned without damaging paint; not fully removable"],
        ["8", "Speed reducer surface", "Both", "Painted (factory finish)",
         "Surface corrosion",
         "Removed as best possible; acid-free vaseline applied"],
        ["9", "Washer\n(motor-to-speed reducer)", "Both", "Carbon steel",
         "Corroded washer",
         "Replaced with SS 316 bolt/washer"],
        ["10", "Cable ties / strips", "Both", "SS 304",
         "Wrong grade \u2014 304 instead of specified 316",
         "All replaced with SS 316 (18 hrs + $525 parts)"],
        ["11", "Edge protection rubber\n(Volvo list)", "Both", "Unknown / non-316",
         "Cannot guarantee offshore suitability",
         "Replaced with confirmed SS 316 quality grey strips"],
        ["12", "Optical switch cable\nouter jacket", "Both", "Polymer jacket",
         "Sharp hex bolts in cable tray cut jacket;\nmoisture ingress path created",
         "Cables repaired; hex bolts addressed"],
        ["13", "End cover flange\n(vent/TW connection)", "Both", "Unpainted bare metal",
         "Missing painting \u2014 bare metal exposed to environment",
         "Flagged in punch list for painting per Norsok M-501"],
    ]

    for row_data in findings:
        row = ft.add_row()
        for j, val in enumerate(row_data):
            set_cell(row.cells[j], val, size=8)

    for row in ft.rows:
        row.cells[0].width = Cm(0.7)
        row.cells[1].width = Cm(3.2)
        row.cells[2].width = Cm(1.3)
        row.cells[3].width = Cm(2.5)
        row.cells[4].width = Cm(4.3)
        row.cells[5].width = Cm(4.8)

    # ── 4. ROOT CAUSE ANALYSIS ──
    doc.add_paragraph()
    add_heading_colored(doc, "4. Root Cause Analysis")

    add_heading_colored(doc, "4.1 Material Specification Gap (Primary Cause)", level=2)
    doc.add_paragraph(
        'The fabricator (TruStop Inc., Phoenix AZ) used SS 304 and generic carbon steel fasteners '
        'where SS 316 was required. Honeywell\'s Bill of Materials (BOM) specified only "SS" '
        '(stainless steel) without grade designation. Sifab explicitly requested SS 316 during FAT '
        'and in the workshop punch list, but this was not reflected in the BOM or procurement.'
    )
    doc.add_paragraph(
        "SS 304 contains less molybdenum than SS 316 and is significantly more susceptible to "
        "chloride-induced pitting and crevice corrosion \u2014 making it unsuitable for North Sea offshore "
        "environments. This single issue accounts for items 6, 9, 10, and 11 in the findings table."
    )

    add_heading_colored(doc, "4.2 Insufficient Surface Protection", level=2)
    doc.add_paragraph(
        "Several surfaces were delivered without adequate coating or painting:"
    )
    for item in [
        "End cover flanges left unpainted (item 13)",
        "Pillow block surfaces with exposed areas (item 2)",
        "Coated bolts with inadequate coating that failed in service (items 4, 5)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "The Norsok M-501 coating specification and Aker BP's additional requirements "
        "(doc 53-000778 Rev 02) were not fully implemented by the factory on ancillary components."
    )

    add_heading_colored(doc, "4.3 Workshop Contamination", level=2)
    doc.add_paragraph(
        "Carbon steel particles from the TruStop workshop contaminated painted surfaces on the motor "
        "and speed reducer (items 7, 8). This is a common fabrication environment issue that was not "
        "adequately controlled. The contamination could not be fully removed without risking damage "
        "to the factory paint finish."
    )

    add_heading_colored(doc, "4.4 Design Gaps for Offshore Service", level=2)
    doc.add_paragraph(
        "Several components were not originally designed or specified for corrosive offshore environments:"
    )
    for item in [
        "Pop rivets on motor compartment covers (carbon steel, item 1)",
        "Seal wires on optical switches (non-316, item 6)",
        "Motor-to-reducer washers (carbon steel, item 9)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "These are standard factory components that TruStop uses for all provers regardless of "
        "operating environment. For North Sea duty, every fastener and ancillary component must be "
        "specified in corrosion-resistant materials."
    )

    # ── 5. COST SUMMARY ──
    add_heading_colored(doc, "5. Remediation Cost Summary")
    doc.add_paragraph(
        "The total cost for all quality concerns (including non-corrosion items such as electrical "
        "rework) was approximately USD 12,021 across 108.5 labour hours plus parts. "
        "The corrosion-specific subset is estimated below:"
    )

    ct = doc.add_table(rows=1, cols=4)
    ct.style = "Table Grid"
    for j, h in enumerate(["Item", "Labour Hours", "Parts (USD)", "Total (USD)"]):
        set_cell(ct.rows[0].cells[j], h, bold=True, size=9, color=WHITE)
    shade_row(ct.rows[0], "003366")

    for row_data in [
        ["Cable ties replacement (304\u2192316)", "18", "525", "2,235"],
        ["Edge protection replacement", "4", "110", "490"],
        ["Optical switch cable repair", "11", "\u2014", "1,192"],
        ["Panel/rivet replacement (Valhall)", "Incl.", "Incl.", "Incl. in labour"],
        ["Pillow block replacement", "Incl.", "Incl.", "Incl. in labour"],
        ["Shock absorber replacement (\u00d72 units)", "Incl.", "Incl.", "Incl. in labour"],
        ["Bolt/washer replacement + preservation", "Incl.", "Incl.", "Incl. in labour"],
        ["TOTAL (corrosion-related subset)", "33+", "635+", "3,917+"],
    ]:
        row = ct.add_row()
        is_total = "TOTAL" in row_data[0]
        for j, val in enumerate(row_data):
            set_cell(row.cells[j], val, bold=is_total, size=9)

    # ── 6. RECOMMENDATIONS ──
    add_heading_colored(doc, "6. Recommendations for Future Projects")
    doc.add_paragraph(
        "Based on the findings from SP-00525/SP-00577, the following recommendations are made "
        "for all future prover deliveries, including the current Snorre A SVP085 project (SP-01415):"
    )

    recs = [
        ("6.1 Explicit SS 316 Specification in BOM",
         'All stainless steel components must be specified as "SS 316" (not generic "SS") in the '
         "Honeywell/TruStop BOM. This includes rivets, cable ties, seal wires, washers, and all "
         "ancillary fasteners. Verify during FAT that material certificates confirm grade 316."),
        ("6.2 Pre-FAT Corrosion Checklist",
         "Develop a checklist of all corrosion-susceptible components to be inspected before FAT. "
         "Include: rivets, bolts, cable ties, seal wires, pillow blocks, shock absorbers, washers, "
         "and all unpainted surfaces. Reject any component not meeting SS 316 or Norsok M-501."),
        ("6.3 Workshop Contamination Control",
         "Require the factory to implement carbon steel separation procedures per Norsok M-501 \u00a75. "
         "No carbon steel work should occur in the same area as prover assembly. Final clean-down "
         "with stainless-safe solvents before packaging."),
        ("6.4 Norsok M-501 Coating \u2014 Full Coverage",
         "Ensure all surfaces are painted per Norsok M-501, including ancillary flanges, pillow "
         "block areas, and end cover connections. Use TSA (System 1) for carbon steel in Zone 1 "
         "and System 6C for flow tube areas. Do NOT use TSA on flow tube assemblies."),
        ("6.5 Preservation Specification",
         "Define preservation requirements in the PO: acid-free vaseline on all machined/threaded "
         "surfaces, VCI packaging for loose fasteners, and humidity indicators in shipping containers."),
        ("6.6 Factory Instruction Package",
         'Issue a written "Factory Instruction Package" to TruStop covering all Norsok and '
         "Aker BP requirements BEFORE manufacturing starts. Include material grades, coating systems, "
         "cable routing requirements, and lessons learned from this project."),
    ]
    for rec_title, rec_body in recs:
        add_heading_colored(doc, rec_title, level=2)
        doc.add_paragraph(rec_body)

    # ── 7. REFERENCES ──
    add_heading_colored(doc, "7. References")
    for ref in [
        "Norsok M-501:2022 \u2014 Surface Preparation and Protective Coating",
        "Aker BP 53-000778 Rev 02 \u2014 Additional Requirements to Norsok M-501",
        "FUI-AKSO-M-SP-00001 Rev 03 \u2014 Coating Specification",
        "Clarification letter: Honeywell NORSOK Coating on Provers",
        "Report: Changing Parts on SVP (Feb 2025)",
        "Prover Quality Concerns \u2014 Factory Response & Discussion (Excel)",
        "Punch List for Hugin A SVP050 (solved with pictures)",
        "Punch List for Valhall SVP015 (solved with pictures)",
        "Comments from End User (Aker BP) \u2014 annotated photos",
        "Mecan Inspection Report (20 November 2024) \u2014 photo documentation",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    # ── FOOTER ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Sifab AS | Bedriftsveien 20, 4313 Sandnes | Org.nr: 886 803 222 | www.sifab.no"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = GREY

    # Save
    output_path = os.path.join(PROJECT_DIR, "Corrosion Analysis Report SP-00525 SP-00577.docx")
    doc.save(output_path)
    print(f"Saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = main()
    os.startfile(path)
