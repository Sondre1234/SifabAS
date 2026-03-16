#!/usr/bin/env python3
"""
Sifab AS — Meeting Recorder & Minutes Generator

Records audio from PC microphone, transcribes with Whisper (local),
and generates meeting minutes in the SIFAB MOM template format.

Usage:
    python tools/meeting_recorder.py record                  # Record audio
    python tools/meeting_recorder.py transcribe <file.wav>   # Transcribe audio
    python tools/meeting_recorder.py minutes <file.wav>      # Full pipeline: transcribe + MOM doc
    python tools/meeting_recorder.py minutes <file.txt>      # Generate MOM from existing transcript
"""

import sys
import os
import json
import datetime
import textwrap
import threading
import queue
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SHARED_DRIVE = Path(os.environ["USERPROFILE"]) / "OneDrive - Sifab AS" / "Dokumenter - Felles"
LOGO_PATH = SHARED_DRIVE / "Logoer" / "SIFAB" / "GB SIFAB LOGO GRØNN 03.jpg"
MOM_TEMPLATE = SHARED_DRIVE / "Maler" / "SIFAB MOM Mal.docx"
OUTPUT_DIR = SHARED_DRIVE / "06 Møtereferat"


# ---------------------------------------------------------------------------
# Recording
# ---------------------------------------------------------------------------
def record_audio(output_path: str | None = None, samplerate: int = 16000) -> str:
    """Record from default microphone until user presses Enter."""
    import sounddevice as sd
    import soundfile as sf

    if output_path is None:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"recording_{ts}.wav"

    print(f"\n  Microphone recording started.")
    print(f"  Press ENTER to stop.\n")

    audio_queue: queue.Queue = queue.Queue()
    recording = True

    def callback(indata, frames, time_info, status):
        if status:
            print(f"  [audio warning: {status}]", file=sys.stderr)
        audio_queue.put(indata.copy())

    # List available devices for debugging
    devices = sd.query_devices()
    default_input = sd.default.device[0]
    if default_input is not None and default_input >= 0:
        dev = devices[default_input]
        print(f"  Using input device: {dev['name']}")
    else:
        print("  Using system default input device")

    stream = sd.InputStream(
        samplerate=samplerate,
        channels=1,
        dtype="float32",
        callback=callback,
    )

    frames_list = []

    def collect():
        while recording:
            try:
                data = audio_queue.get(timeout=0.5)
                frames_list.append(data)
            except queue.Empty:
                continue

    collector = threading.Thread(target=collect, daemon=True)

    with stream:
        collector.start()
        try:
            input()  # blocks until Enter
        except (KeyboardInterrupt, EOFError):
            pass
        recording = False
        collector.join(timeout=2)

    if not frames_list:
        print("  No audio captured.")
        sys.exit(1)

    import numpy as np
    audio_data = np.concatenate(frames_list, axis=0)
    duration = len(audio_data) / samplerate

    sf.write(output_path, audio_data, samplerate)
    print(f"  Saved {duration:.1f}s of audio to: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Transcription (local Whisper via faster-whisper)
# ---------------------------------------------------------------------------
def transcribe_audio(audio_path: str, model_size: str = "base") -> str:
    """Transcribe audio file using faster-whisper (local, no API needed)."""
    from faster_whisper import WhisperModel

    print(f"\n  Loading Whisper model '{model_size}'...")
    model = WhisperModel(model_size, device="cpu", compute_type="int8")

    print(f"  Transcribing: {audio_path}")
    segments, info = model.transcribe(
        audio_path,
        language="no",  # Norwegian
        beam_size=5,
        vad_filter=True,
    )

    print(f"  Detected language: {info.language} (probability {info.language_probability:.1%})")

    lines = []
    for segment in segments:
        ts = f"[{_fmt_time(segment.start)} - {_fmt_time(segment.end)}]"
        lines.append(f"{ts}  {segment.text.strip()}")
        print(f"  {ts}  {segment.text.strip()}")

    transcript = "\n".join(lines)

    # Save transcript alongside audio
    txt_path = Path(audio_path).with_suffix(".txt")
    txt_path.write_text(transcript, encoding="utf-8")
    print(f"\n  Transcript saved to: {txt_path}")

    return transcript


def _fmt_time(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


# ---------------------------------------------------------------------------
# Meeting minutes generation (Claude API)
# ---------------------------------------------------------------------------
MINUTES_PROMPT = """\
You are a professional meeting minutes writer for Sifab AS, a Norwegian energy company.

Given the following transcript of a meeting (spoken in Norwegian), produce structured meeting minutes in ENGLISH.

Output a JSON object with these fields:
- "title": short meeting title (English)
- "purpose": purpose of the meeting (English, 1-2 sentences)
- "attendees": list of attendee names (extract from transcript if mentioned, otherwise ["To be confirmed"])
- "items": list of objects, each with:
  - "number": item number (1, 2, 3...)
  - "discussion": summary of what was discussed (English, concise but complete)
  - "action": who needs to do what and by when (English), or "N/A" if no action

Be concise but capture all decisions, action items, and key discussion points.
Translate everything to English. Keep technical terms (product names, project numbers) as-is.

TRANSCRIPT:
{transcript}
"""


def generate_minutes_data(transcript: str) -> dict:
    """Use Claude API to generate structured meeting minutes from transcript."""
    try:
        import anthropic

        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            print("  No ANTHROPIC_API_KEY found. Generating basic minutes without AI summarization.")
            return _fallback_minutes(transcript)

        client = anthropic.Anthropic(api_key=api_key)

        print("  Generating meeting minutes with Claude...")
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": MINUTES_PROMPT.format(transcript=transcript),
            }],
        )

        text = response.content[0].text
        # Extract JSON from response
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]

        return json.loads(text.strip())

    except ImportError:
        print("  anthropic package not installed. Using fallback.")
        return _fallback_minutes(transcript)
    except Exception as e:
        print(f"  Claude API error: {e}")
        print("  Using fallback minutes format.")
        return _fallback_minutes(transcript)


def _fallback_minutes(transcript: str) -> dict:
    """Generate basic minutes structure without AI."""
    return {
        "title": "Meeting Minutes",
        "purpose": "See transcript below",
        "attendees": ["To be confirmed"],
        "items": [{
            "number": 1,
            "discussion": transcript[:2000] + ("..." if len(transcript) > 2000 else ""),
            "action": "Review and update",
        }],
    }


# ---------------------------------------------------------------------------
# MOM Document generation (.docx)
# ---------------------------------------------------------------------------
def create_mom_docx(minutes_data: dict, output_path: str | None = None, transcript: str = "") -> str:
    """Create a SIFAB MOM .docx from structured minutes data."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    today = datetime.date.today().strftime("%d.%m.%Y")
    title = minutes_data.get("title", "Meeting Minutes")
    purpose = minutes_data.get("purpose", "")
    attendees = minutes_data.get("attendees", [])
    items = minutes_data.get("items", [])

    # ── Header table (mimics SIFAB MOM template) ──
    header_table = doc.add_table(rows=6, cols=4)
    header_table.style = "Table Grid"

    # Row 0: Logo + MINUTES OF MEETING
    row0 = header_table.rows[0]
    # Try to add logo
    if LOGO_PATH.exists():
        try:
            row0.cells[0].paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Cm(3))
        except Exception:
            _set_cell(row0.cells[0], "SIFAB AS", bold=True, size=14)
    else:
        _set_cell(row0.cells[0], "SIFAB AS", bold=True, size=14)

    _merge_cells(row0.cells[1], row0.cells[3])
    _set_cell(row0.cells[1], "MINUTES OF MEETING", bold=True, size=16, align="center",
              color=RGBColor(0, 51, 102))

    # Row 1: Location / Issued By / Meeting Date / Date of Issue
    row1 = header_table.rows[1]
    _set_cell(row1.cells[0], "LOCATION OF MEETING\nSifab AS, Sandnes", size=9)
    _set_cell(row1.cells[1], "ISSUED BY\nSondre Falch", size=9)
    _set_cell(row1.cells[2], f"MEETING DATE\n{today}", size=9)
    _set_cell(row1.cells[3], f"DATE OF ISSUE\n{today}", size=9)

    # Row 2: Title / Contract No.
    row2 = header_table.rows[2]
    _merge_cells(row2.cells[0], row2.cells[2])
    _set_cell(row2.cells[0], f"TITLE:\n{title}", size=9)
    _set_cell(row2.cells[3], "CONTRACT NO.\n", size=9)

    # Row 3: Purpose / Reference No.
    row3 = header_table.rows[3]
    _merge_cells(row3.cells[0], row3.cells[2])
    _set_cell(row3.cells[0], f"PURPOSE OF MEETING:\n{purpose}", size=9)
    _set_cell(row3.cells[3], "REFERENCE NO.\n", size=9)

    # Row 4: Attended By / Distribution
    row4 = header_table.rows[4]
    _merge_cells(row4.cells[0], row4.cells[2])
    attendees_str = "\n".join(f"  - {a}" for a in attendees) if attendees else "  To be confirmed"
    _set_cell(row4.cells[0], f"ATTENDED BY:\n{attendees_str}", size=9)
    _set_cell(row4.cells[3], "DISTRIBUTION:\nAll Attendees", size=9)

    # Row 5: separator
    row5 = header_table.rows[5]
    for cell in row5.cells:
        cell.text = ""

    doc.add_paragraph()  # spacer

    # ── Content table ──
    n_items = max(len(items), 1)
    content_table = doc.add_table(rows=1 + n_items, cols=3)
    content_table.style = "Table Grid"

    # Set column widths
    for row in content_table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(12)
        row.cells[2].width = Cm(3.5)

    # Header row
    hdr = content_table.rows[0]
    _set_cell(hdr.cells[0], "ITEM", bold=True, size=9, align="center")
    _set_cell(hdr.cells[1], "DISCUSSIONS", bold=True, size=9, align="center")
    _set_cell(hdr.cells[2], "ACTION /\nDATE / BY", bold=True, size=9, align="center")

    # Shade header
    for cell in hdr.cells:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "003366",
        })
        shading.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Content rows
    for i, item in enumerate(items):
        row = content_table.rows[1 + i]
        _set_cell(row.cells[0], str(item.get("number", i + 1)), size=9, align="center")
        _set_cell(row.cells[1], item.get("discussion", ""), size=9)
        _set_cell(row.cells[2], item.get("action", "N/A"), size=9)

    doc.add_paragraph()  # spacer

    # ── Appendix: Full Transcript ──
    if transcript:
        doc.add_paragraph()
        app_heading = doc.add_paragraph()
        run = app_heading.add_run("APPENDIX — FULL TRANSCRIPT (Norwegian)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

        tp = doc.add_paragraph()
        run = tp.add_run(transcript)
        run.font.size = Pt(8)
        run.font.name = "Consolas"

    # ── Footer ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "Sifab AS | Bedriftsveien 20, 4313 Sandnes | Org.nr: 886 803 222 | www.sifab.no"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    if output_path is None:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title)[:50].strip()
        output_path = str(OUTPUT_DIR / f"MOM {today} - {safe_title}.docx")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"\n  Meeting minutes saved to:\n  {output_path}")
    return output_path


def _set_cell(cell, text, bold=False, size=10, align="left", color=None):
    """Set cell text with formatting."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _merge_cells(cell_start, cell_end):
    """Merge table cells."""
    cell_start.merge(cell_end)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)

    command = sys.argv[1].lower()

    if command == "record":
        output = sys.argv[2] if len(sys.argv) > 2 else None
        record_audio(output)

    elif command == "transcribe":
        if len(sys.argv) < 3:
            print("  Usage: meeting_recorder.py transcribe <audio_file.wav>")
            sys.exit(1)
        audio_path = sys.argv[2]
        model_size = sys.argv[3] if len(sys.argv) > 3 else "base"
        transcribe_audio(audio_path, model_size)

    elif command == "minutes":
        if len(sys.argv) < 3:
            print("  Usage: meeting_recorder.py minutes <file.wav|file.txt>")
            sys.exit(1)

        input_path = sys.argv[2]
        ext = Path(input_path).suffix.lower()

        if ext == ".txt":
            # Already a transcript
            transcript = Path(input_path).read_text(encoding="utf-8")
            print(f"  Loaded transcript from: {input_path}")
        else:
            # Audio file — transcribe first
            model_size = sys.argv[3] if len(sys.argv) > 3 else "base"
            transcript = transcribe_audio(input_path, model_size)

        # Generate structured minutes
        minutes_data = generate_minutes_data(transcript)

        # Create MOM document
        create_mom_docx(minutes_data, transcript=transcript)

    elif command == "full":
        # Full pipeline: record → transcribe → minutes
        output_wav = sys.argv[2] if len(sys.argv) > 2 else None
        model_size = sys.argv[3] if len(sys.argv) > 3 else "base"

        wav_path = record_audio(output_wav)
        transcript = transcribe_audio(wav_path, model_size)
        minutes_data = generate_minutes_data(transcript)
        mom_path = create_mom_docx(minutes_data, transcript=transcript)

        # Open the document
        print(f"\n  Opening document...")
        os.startfile(mom_path)

    else:
        print(f"  Unknown command: {command}")
        print(__doc__)
        sys.exit(1)


if __name__ == "__main__":
    main()
