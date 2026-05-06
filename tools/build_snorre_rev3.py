import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# LOGO
logo_path = Path('C:/Users/SondreFalch/SifabAS/temp_logo.jpeg')
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = logo_para.add_run()
run.add_picture(str(logo_path), width=Cm(6))

# TITLE
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SVP085 MODULAR SPLIT EXECUTION PLAN')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 80, 40)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Snorre A Platform \u2014 Small Volume Prover')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# DOCUMENT INFO TABLE
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ('Project', 'SP-01415 \u2014 Small Volume Prover Snorre A'),
    ('RFQ Reference', 'GM-8501-1447 (Equinor via Guidant)'),
    ('Equipment', 'Honeywell Calibron SVP085'),
    ('Revision', 'Rev 3 \u2014 19 March 2026'),
    ('Prepared by', 'Sifab AS (Tom S. Falch / Sondre Falch)'),
    ('Status', 'FOR REVIEW \u2014 Guidant / Equinor'),
]
for i, (label, value) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph()

# REVISION HISTORY
doc.add_heading('Revision History', level=2)
rev_table = doc.add_table(rows=4, cols=3)
rev_table.style = 'Light Grid Accent 1'
rev_headers = ['Rev', 'Date', 'Description']
for i, hdr in enumerate(rev_headers):
    rev_table.rows[0].cells[i].text = hdr
    for p in rev_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

rev_data = [
    ('1', '26 Feb 2026', 'Initial draft \u2014 3-section frame split concept, internal review'),
    ('2', '27 Feb 2026', 'Updated mounting pad terminology, top-mounted flanges note, Guidant access photos'),
    ('3', '19 Mar 2026', 'Revised to 2-section frame split per Honeywell/Guidant feedback. Transport module list updated. Painting strategy confirmed. Guidant and Equinor approval received.'),
]
for r_idx, (rev, date, desc) in enumerate(rev_data):
    rev_table.rows[r_idx + 1].cells[0].text = rev
    rev_table.rows[r_idx + 1].cells[1].text = date
    rev_table.rows[r_idx + 1].cells[2].text = desc

doc.add_page_break()

# ============ 1. BACKGROUND ============
doc.add_heading('1. Background', level=1)
doc.add_paragraph(
    'Equinor has issued RFQ GM-8501-1447 for a Honeywell Calibron SVP085 Small Volume Prover '
    'to be installed on the Snorre A platform in the North Sea. Guidant Measurement is the EPC '
    'contractor responsible for the metering package, and Sifab AS is the authorised Honeywell channel '
    'partner providing the prover.'
)
doc.add_paragraph(
    'The SVP085 must be transported through the platform access door and along an internal corridor '
    'to reach the installation location. The prover does not fit through the access opening as a '
    'complete unit \u2014 the frame width of 1,448 mm exceeds practical transport limits when accounting '
    'for the corridor turns.'
)
doc.add_paragraph(
    'This document describes the modular split concept that enables transport and installation of the '
    'SVP085 on Snorre A, developed in consultation with Honeywell Enraf, Guidant, and Equinor.'
)

# ============ 2. PLATFORM ACCESS ============
doc.add_heading('2. Platform Access Constraints', level=1)

dim_table = doc.add_table(rows=2, cols=3)
dim_table.style = 'Light Grid Accent 1'
dim_table.rows[0].cells[0].text = ''
dim_table.rows[0].cells[1].text = 'SVP085'
dim_table.rows[0].cells[2].text = 'Door / Corridor Limit'
for c in [1, 2]:
    for p in dim_table.rows[0].cells[c].paragraphs:
        for r in p.runs:
            r.bold = True
dim_table.rows[1].cells[0].text = 'Prover Length'
dim_table.rows[1].cells[1].text = '5,258 mm'
dim_table.rows[1].cells[2].text = 'Max ~2,450 mm (corridor)'
for p in dim_table.rows[1].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True

doc.add_paragraph()

doc.add_heading('Platform Access Confirmed', level=2)
doc.add_paragraph(
    'Guidant (Torleif Espegard, Senior Tender Engineer) has confirmed the following based on site '
    'photos and customer review (March 2026):'
)
for b in [
    'Max 1.59 m width can work if ladder east of metering package is temporarily removed (currently not required)',
    'Transport route confirmed: through door, along corridor, turn west to installation area',
    'Equinor (customer) has reviewed the split concept and confirmed it should work fine',
    'Guidant approved the proposed supervision time estimate for Honeywell/Intertek activities',
]:
    doc.add_paragraph(b, style='List Bullet')

# ============ 3. SPLIT CONCEPT ============
doc.add_heading('3. Split Concept \u2014 2-Section Frame', level=1)

doc.add_paragraph(
    'Based on input from Honeywell (Samir Sakota, Mechanical Engineer) and Guidant feedback, '
    'the frame will be split into 2 sections with a bolted splice joint at approximately the midpoint. '
    'This is the simplest and lowest-risk approach.'
)

doc.add_heading('Frame Split Details', level=2)
for b in [
    'Total frame length: 5,258 mm',
    'Split into 2 sections of approximately 2,630 mm each',
    'Sections transported on their side (tilted 90\u00b0) \u2014 effective width becomes ~200\u2013300 mm instead of 1,448 mm lying flat',
    'Sections bolted together in field at installation location on Snorre A',
    'Existing mounting pads at each end of the frame carry the main structural loads',
    'Additional mounting pad(s) at the splice joint for support \u2014 design by Honeywell',
    'Drive end mounting pad must be heavy-duty to absorb piston cycling reaction forces from the motor',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Flow Tube', level=2)
doc.add_paragraph('The flow tube is removed from the frame and transported separately:')
for b in [
    'Flow tube body length (without end flanges): approximately 2,600\u20132,700 mm',
    'Flow tube weight: approximately 3,500 kg',
    'Width: approximately 1,000 mm \u2014 fits through 1,400 mm door',
    'This SVP085 configuration has inlet/outlet flanges mounted on top, increasing overall height beyond the standard 1,316 mm manual specification',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Lifting approach: Two 12" CL600 blind flanges with welded lifting ears will be installed on the '
    'top-mounted inlet/outlet flange connections. This provides a stable lifting base when the flow tube '
    'is being positioned onto the reassembled frame on Snorre A. Lifting analysis to be verified.'
)

# ============ 4. TRANSPORT MODULES ============
doc.add_heading('4. Transport Modules', level=1)
doc.add_paragraph(
    "The prover will be disassembled at the onshore fabrication facility in Norway after Guidant's "
    'flow test is completed, and packed into the following transport modules:'
)

mod_table = doc.add_table(rows=8, cols=3)
mod_table.style = 'Light Grid Accent 1'
for i, hdr in enumerate(['Module', 'Description', 'Approx. Dimensions / Weight']):
    mod_table.rows[0].cells[i].text = hdr
    for p in mod_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

modules = [
    ('1', 'Frame Section A (drive end)', '~2,630 x 300 x 1,448 mm (on side)'),
    ('2', 'Frame Section B (non-drive end)', '~2,630 x 300 x 1,448 mm (on side)'),
    ('3', 'Flow Tube (in wooden crate)', '~2,900 x 600 mm, ~3,500 kg'),
    ('4', 'End Flanges + Piston Assembly', 'Wooden crate, protected RTJ faces'),
    ('5', 'Drive System (removed as one unit)', 'Wooden crate'),
    ('6', 'Electrical Components + Controller', 'Wooden crate'),
    ('7', 'Assembly Materials + Hardware', 'Wooden crate'),
]
for r_idx, (mod, desc, dim) in enumerate(modules):
    mod_table.rows[r_idx + 1].cells[0].text = mod
    mod_table.rows[r_idx + 1].cells[1].text = desc
    mod_table.rows[r_idx + 1].cells[2].text = dim

doc.add_paragraph()
doc.add_paragraph(
    'A transport trolley with wheels will be used to move modules through the platform corridor. '
    'Exact design to be coordinated with Guidant and platform logistics team \u2014 existing platform '
    'equipment may be suitable given the routine handling of large items on Snorre A.'
)

# ============ 5. EXECUTION SEQUENCE ============
doc.add_heading('5. Execution Sequence', level=1)

phases = [
    ('Phase 1: Factory Assembly & Testing (Honeywell)', [
        'SVP085 manufactured and assembled at Honeywell facility',
        'Frame designed from the start with bolted splice joint (design-for-split)',
        'Factory calibration and acceptance testing as complete unit',
        'Gravimetric calibration certificate issued',
    ]),
    ('Phase 2: Shipping to Norway & Flow Test (Guidant)', [
        'Complete prover shipped to Guidant fabrication location in Norway',
        'Prover used by Guidant for ultrasonic meter testing on the metering package',
        'Flow test completed as a fully assembled unit',
    ]),
    ('Phase 3: Disassembly & Painting (Sifab)', [
        'After flow test, prover disassembled into transport modules at onshore facility',
        'Flow tube packed in wooden crate and transported to painting workshop',
        'Flow tube and end flanges painted after flow test (not before)',
        'Frame sections and other modules packed for offshore transport',
    ]),
    ('Phase 4: Transport to Snorre A', [
        'Transport modules shipped to Snorre A platform',
        'Each module brought through access door and along corridor to installation area',
        'Frame sections tilted 90\u00b0 on transport trolley for corridor passage',
        'Flow tube crate rolled through on transport wheels (only ~600 mm wide)',
    ]),
    ('Phase 5: Reassembly on Snorre A (Sifab + Honeywell)', [
        'Frame sections positioned and bolted together at splice joint',
        'Mounting pads aligned and secured to deck',
        'Flow tube lifted using blind flange lifting ears and mounted on reassembled frame',
        'Drive system reconnected to flow tube',
        'Electrical components and controller installed and connected',
        'Piping and process connections completed',
        'Honeywell Enraf supervisor on site for reassembly verification',
        'Intertek supporting offshore work (Equinor frame agreement holder for offshore proving services)',
        'Blu Electro for electrical work offshore',
    ]),
    ('Phase 6: Commissioning & Acceptance', [
        'System leak test and functional checks',
        'Prover verification run (water draw or master meter comparison)',
        'Honeywell sign-off confirming reassembly meets factory specifications',
        'Warranty starts from date Honeywell has fully assembled and released SVP for use on site',
    ]),
]

for phase_title, bullets in phases:
    doc.add_heading(phase_title, level=2)
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

# ============ 6. KEY ENGINEERING POINTS ============
doc.add_heading('6. Key Engineering Points', level=1)
for b in [
    'Engineering of prover frame split by Honeywell. Reassembly guidelines and procedures to be provided by Honeywell.',
    'Frame designed with bolted splice joint from the start \u2014 no field welding planned.',
    'Honeywell Enraf supervisor required on site during reassembly to ensure compliance with manufacturer guidelines.',
    'Intertek engaged for offshore proving services support (Equinor frame agreement holder).',
    'Blu Electro engaged for offshore electrical work.',
    'Flow tube 0.020% repeatability must be maintained after disassembly and reassembly \u2014 to be confirmed and warranted by Honeywell.',
    'All RTJ flange faces must be protected during transport. Piston seal inspection required before reinstallation.',
    'Warranty shall be valid from the date that Honeywell has fully assembled and released the SVP for use on site.',
]:
    doc.add_paragraph(b, style='List Bullet')

# ============ 7. OPEN ITEMS ============
doc.add_heading('7. Open Items', level=1)
doc.add_paragraph('The following items require confirmation or input:')

tq_table = doc.add_table(rows=7, cols=3)
tq_table.style = 'Light Grid Accent 1'
for i, hdr in enumerate(['#', 'Item', 'Status']):
    tq_table.rows[0].cells[i].text = hdr
    for p in tq_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

tq_data = [
    ('1', 'Honeywell BOM and detailed splice joint design', 'In progress (Samir Sakota)'),
    ('2', 'Exact flow tube height with top-mounted 12" CL600 RTJ flanges', 'Awaiting Honeywell'),
    ('3', 'Transport trolley design for platform corridor', 'To be coordinated with Guidant'),
    ('4', 'Lifting analysis for blind flange lifting ear concept (~3,500 kg)', 'Sifab to verify'),
    ('5', 'Painting specification for flow tube and end flanges', 'TBD after flow test'),
    ('6', 'Offshore reassembly man-hours estimate (Honeywell + Intertek + Blu Electro)', 'In progress'),
]
for r_idx, (num, item, status) in enumerate(tq_data):
    tq_table.rows[r_idx + 1].cells[0].text = num
    tq_table.rows[r_idx + 1].cells[1].text = item
    tq_table.rows[r_idx + 1].cells[2].text = status

doc.add_paragraph()

# ============ 8. CONTACTS ============
doc.add_heading('8. Project Contacts', level=1)

ct = doc.add_table(rows=7, cols=4)
ct.style = 'Light Grid Accent 1'
for i, hdr in enumerate(['Name', 'Company', 'Role', 'Contact']):
    ct.rows[0].cells[i].text = hdr
    for p in ct.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

contacts = [
    ('Tom Sverre Falch', 'Sifab AS', 'General Manager', 'tom.falch@sifab.no | +47 416 284 08'),
    ('Sondre Falch', 'Sifab AS', 'Sales Manager', 'sondre.falch@sifab.no | +47 900 29 588'),
    ('Sidney Swart', 'Honeywell Enraf', 'Sales', 'Sidney.Swart@Honeywell.com'),
    ('Samir Sakota', 'Honeywell Enraf', 'Mechanical Engineer', 'Samir.Sakota@Honeywell.com'),
    ('Torleif Espegard', 'Guidant Measurement', 'Sr. Tender Engineer', 'torleif.espegard@guidantmeasurement.com'),
    ('Marco Rosas', 'Honeywell Enraf', 'Engineering', 'Marco.Rosas@Honeywell.com'),
]
for r_idx, (name, company, role, contact) in enumerate(contacts):
    ct.rows[r_idx + 1].cells[0].text = name
    ct.rows[r_idx + 1].cells[1].text = company
    ct.rows[r_idx + 1].cells[2].text = role
    ct.rows[r_idx + 1].cells[3].text = contact

doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Sifab AS | Bedriftsveien 20, 4313 Sandnes, Norway | www.sifab.no')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

# SAVE
shared = Path(os.environ['USERPROFILE']) / 'OneDrive - Sifab AS' / 'Dokumenter - Felles'
output_path = shared / 'Zigma360' / 'Projects' / 'SP-01415 Small Volume Prover Snorre A' / 'Modular_Split_Execution_Plan_SVP085_Rev3.docx'
doc.save(str(output_path))
print(f'Document saved: {output_path}')
