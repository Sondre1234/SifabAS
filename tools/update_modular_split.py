import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc_path = 'C:/Users/SondreFalch/SifabAS/projects/snorre-a-compact-prover/engineering/Modular_Split_Execution_Plan_SVP085.docx'
doc = Document(doc_path)

# ============================================================
# STEP 1: Update the header table - change date and status
# ============================================================
header_table = doc.tables[0]
for row in header_table.rows:
    for cell in row.cells:
        if '2026-02-26' in cell.text:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '2026-02-26' in run.text:
                        run.text = run.text.replace('2026-02-26', '2026-02-27')
        if 'DRAFT' in cell.text:
            for para in cell.paragraphs:
                for run in para.runs:
                    if 'DRAFT' in run.text:
                        run.text = run.text.replace(
                            'DRAFT',
                            'REVISION 1'
                        )
                    if 'For internal review and Honeywell discussion' in run.text:
                        run.text = run.text.replace(
                            'For internal review and Honeywell discussion',
                            'Updated with Guidant access photos and GA drawing split illustration'
                        )

# ============================================================
# STEP 2: Replace 2,500 mm with 2,450 mm throughout
# ============================================================
replacements_done = 0
for para in doc.paragraphs:
    for run in para.runs:
        if '2,500' in run.text:
            run.text = run.text.replace('2,500', '2,450')
            replacements_done += 1
        if '2500' in run.text:
            run.text = run.text.replace('2500', '2450')
            replacements_done += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '2,500' in run.text:
                        run.text = run.text.replace('2,500', '2,450')
                        replacements_done += 1
                    if '2500' in run.text:
                        run.text = run.text.replace('2500', '2450')
                        replacements_done += 1

print(f"Replaced 2500->2450 in {replacements_done} locations")

# ============================================================
# Helper functions
# ============================================================

def add_heading_before(ref_el, text, level=2):
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    new_p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    ref_el.addprevious(new_p)
    return new_p

def add_text_before(ref_el, text, bold=False, italic=False):
    new_p = OxmlElement('w:p')
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i_el = OxmlElement('w:i')
            rPr.append(i_el)
        if bold or italic:
            r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
    ref_el.addprevious(new_p)
    return new_p

def add_image_before(ref_el, doc, image_path, width_inches=5.5):
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_el.addprevious(temp_para._element)
    return temp_para

def add_separator_before(ref_el):
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '____________________________________________________________'
    r.append(t)
    new_p.append(r)
    ref_el.addprevious(new_p)
    return new_p

# ============================================================
# STEP 3: Insert new section 1b after Problem Statement
# ============================================================

ref_data_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '2. Reference Data':
        ref_data_idx = i
        break

print(f"Reference Data heading at paragraph {ref_data_idx}")
ref_element = doc.paragraphs[ref_data_idx]._element

add_separator_before(ref_element)

add_heading_before(ref_element, '1b. Platform Access Confirmation (Guidant/Equinor, 27 Feb 2026)', level=2)

add_text_before(ref_element,
    'On 27 February 2026, Guidant (Torleif Espegard, Senior Tender Engineer) provided site photos and access confirmation for the Snorre A platform corridor. Key findings:')

add_text_before(ref_element, '')

add_text_before(ref_element,
    'Welding OK: Equinor confirms welding is acceptable on the platform. Preferred before prover is in position (unless for supplementary frame). Design drawing of welding scope requested.', bold=True)

add_text_before(ref_element, '')

add_text_before(ref_element,
    'Width Clearance: The wooden frame template is 1.4 m wide. Guidant/Equinor believe 1.59 m width can work if the ladder east of the existing metering package is temporarily removed. This means frame sections (1,448 mm when flipped) should fit.', bold=True)

add_text_before(ref_element, '')

add_text_before(ref_element,
    'Customer Question: "What do you envision for length and height?" Our answer: Max 2,450 mm length per module, max ~1,448 mm height when frame sections are flipped 90 degrees on their side.', bold=True)

add_text_before(ref_element, '')

add_text_before(ref_element,
    'Routing Confirmed: Transport route from access door through corridor toward existing metering package skid, with a turn westward to prover installation location.', bold=True)

add_text_before(ref_element, '')

add_heading_before(ref_element, 'Snorre A Corridor Access Photos (Guidant, 27 Feb 2026)', level=3)

add_text_before(ref_element,
    'Wooden frame templates placed on deck to verify physical clearance for prover module transport:')

add_text_before(ref_element, '')

add_text_before(ref_element, 'Photo 1: Corridor view with existing metering package skid and wooden size template', bold=True)
add_image_before(ref_element, doc, 'C:/Users/SondreFalch/SifabAS/downloads/image001.png', width_inches=4.5)
add_text_before(ref_element, '')

add_text_before(ref_element, 'Photo 2: Access door area with wooden frame (1.4m wide, 1.59m possible with ladder removal)', bold=True)
add_image_before(ref_element, doc, 'C:/Users/SondreFalch/SifabAS/downloads/image002.png', width_inches=4.0)
add_text_before(ref_element, '')

add_text_before(ref_element, 'Photo 3: Corridor length view with access door on right', bold=True)
add_image_before(ref_element, doc, 'C:/Users/SondreFalch/SifabAS/downloads/image003.png', width_inches=4.5)
add_text_before(ref_element, '')

add_text_before(ref_element, 'Photo 4: Transport routing toward existing metering package skid, turning prover westward', bold=True)
add_image_before(ref_element, doc, 'C:/Users/SondreFalch/SifabAS/downloads/image004.png', width_inches=4.0)
add_text_before(ref_element, '')

print("Added section 1b with Guidant corridor photos")

# ============================================================
# STEP 4: Insert split illustration in Section 4
# ============================================================

split_concept_idx = None
for i, para in enumerate(doc.paragraphs):
    if 'Proposed 3-Module Split Concept' in para.text:
        split_concept_idx = i
        break

module1_idx = None
for i, para in enumerate(doc.paragraphs):
    if 'Module 1:' in para.text and (split_concept_idx is None or i > split_concept_idx):
        module1_idx = i
        break

print(f"Module 1 heading at paragraph index (approx): searching...")

if module1_idx:
    ref_el = doc.paragraphs[module1_idx]._element

    add_text_before(ref_el, '')

    add_text_before(ref_el,
        'Figure: SVP085 Modular Split Illustration (annotated from Honeywell GA Drawing, Manual Part No. 44200001)', bold=True)
    add_image_before(ref_el, doc, 'C:/Users/SondreFalch/SifabAS/downloads/SVP085_split_illustration.png', width_inches=6.0)
    add_text_before(ref_el,
        'Red dashed lines = bolted splice joint locations. Green arrows = additional saddle-type mounting brackets at each splice point. Frame splits into 3 sections of approx. 1,750 mm each (all within 2,450 mm max). Flow tube (2,954 mm bore) transported separately as Module 1 on temporary frame.')

    add_text_before(ref_el, '')

    add_text_before(ref_el,
        'Figure: SVP085 Prover Critical Dimensions (Honeywell Manual Part No. 44200001, Table 21)', bold=True)
    add_image_before(ref_el, doc, 'C:/Users/SondreFalch/SifabAS/downloads/prover_manual_images/page32_full.png', width_inches=5.5)

    add_text_before(ref_el, '')

    add_text_before(ref_el,
        'Figure: SVP085 Anchor Points on Frame (Figure 2-4 from Manual)', bold=True)
    add_image_before(ref_el, doc, 'C:/Users/SondreFalch/SifabAS/downloads/prover_manual_images/page39_full.png', width_inches=5.0)

    add_text_before(ref_el, '')

    print("Added split illustrations to Section 4")

# ============================================================
# STEP 5: Update footer address
# ============================================================
for para in doc.paragraphs:
    for run in para.runs:
        if 'Bedriftsveien 24' in run.text:
            run.text = run.text.replace('Bedriftsveien 24', 'Bedriftsveien 20')

# ============================================================
# SAVE
# ============================================================
output_path = 'C:/Users/SondreFalch/SifabAS/projects/snorre-a-compact-prover/engineering/Modular_Split_Execution_Plan_SVP085_Rev1.docx'
doc.save(output_path)
print(f"\nSaved: {output_path}")
print("Done!")
