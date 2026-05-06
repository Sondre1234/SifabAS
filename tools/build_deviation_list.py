"""Generate SP-01415 Deviation List and Compliance Matrix as Word documents."""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SHARED = Path(os.environ['USERPROFILE']) / 'OneDrive - Sifab AS' / 'Dokumenter - Felles'
DEST = SHARED / 'Zigma360' / 'Projects' / 'SP-01415 Small Volume Prover Snorre A' / '05 Dokumentasjon' / '04.Dok sendt til kunde'


def set_cell(cell, text, bold=False, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold


def add_header_table(doc, pairs):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(pairs):
        set_cell(t.rows[i].cells[0], k, bold=True)
        set_cell(t.rows[i].cells[1], v)
    return t


def severity_color(sev):
    if sev == 'HIGH':
        return RGBColor(204, 0, 0)
    elif sev == 'MEDIUM':
        return RGBColor(204, 102, 0)
    return RGBColor(0, 128, 0)


DEVIATIONS = [
    {
        'id': 'DEV-001', 'title': 'Ingress Protection - IP65/IP56 vs. IP66',
        'severity': 'HIGH',
        'rfq_ref': 'S1.2 Area Classification',
        'requirement': 'Ingress Protection: min. IP66',
        'offered': 'IP65 for transmitters and junction boxes; IP56 for motor',
        'deviation': 'IP65 and IP56 are both below the required IP66. IP65 lacks protection against powerful water jets; IP56 lacks full dust protection.',
        'risk_owner': 'Honeywell (prover manufacturer)',
        'resolution': 'Honeywell to confirm if IP66-rated alternatives are available for motor and junction boxes, with any price/schedule impact. If not possible, a formal deviation request must be submitted to Guidant/Equinor for acceptance. Sifab does not accept risk on this item.',
    },
    {
        'id': 'DEV-002', 'title': 'Motor Frequency - 50 Hz vs. 60 Hz',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.5 Motor',
        'requirement': '230 VAC, 3-phase, 60 Hz',
        'offered': 'Initially quoted 50 Hz. Sidney confirmed 60 Hz possible with no USD impact (23 Mar).',
        'deviation': 'Quote and technical spec still show 50 Hz. Needs formal update.',
        'risk_owner': 'Honeywell',
        'resolution': 'Honeywell to issue updated proposal confirming 60 Hz motor with final pricing. Sifab to pass through any adder to Guidant.',
    },
    {
        'id': 'DEV-003', 'title': 'Welding Standards - ASME IX vs. NORSOK/TR',
        'severity': 'HIGH',
        'rfq_ref': 'S1.11 Fabrication',
        'requirement': 'Welding per TR1826, NORSOK M-601, NORSOK L-004.',
        'offered': 'Pressure containing welds per ASME BPV code section IX.',
        'deviation': 'Honeywell welds per ASME IX, not NORSOK M-601 / TR1826. Same approach used on Aker BP provers via project deviation request.',
        'risk_owner': 'To be accepted by Guidant/Equinor via formal deviation request',
        'resolution': 'Submit project deviation request to Guidant for flow tube welding, referencing Aker BP precedent. Honeywell to provide ASME IX WPS/WPQ documentation. Sifab does not carry welding risk.',
    },
    {
        'id': 'DEV-004', 'title': 'NORSOK Painting - Excluded from Honeywell Scope',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.9 Painting',
        'requirement': 'All surface prep per TR0042 and NORSOK M-501. Flow tube: System 6C.',
        'offered': 'NORSOK Painting not included. Sifab to source locally.',
        'deviation': 'Honeywell delivers prover without NORSOK M-501 coating. Motor/gearbox with Honeywell offshore paint - exemption requested.',
        'risk_owner': 'Sifab (local painting), Guidant/Equinor (motor/gearbox exemption)',
        'resolution': 'Sifab to arrange NORSOK M-501 System 6C painting in Norway after flow test. Motor/gearbox: Honeywell to provide certificates. All parts under insulation to be painted (confirmed 23 Mar).',
    },
    {
        'id': 'DEV-005', 'title': 'Warranty - 24 Months Delivery vs. 28 Months Installation',
        'severity': 'HIGH',
        'rfq_ref': 'S1 Scope of Supply',
        'requirement': 'Warranty from date vendor has assembled and released SVP on site. Min. 28 months after installation.',
        'offered': 'Standard 24 months after delivery. Parts only, excludes wearable parts and labor.',
        'deviation': 'Honeywell warranty starts at delivery (not installation), is 24 months (not 28), and excludes labor and wearables.',
        'risk_owner': 'Sifab exposed to gap between HW 24m/delivery and Guidant 28m/installation',
        'resolution': 'Formalize extended warranty with Honeywell (EWA option). Back-to-back warranty terms required. Sifab must not accept warranty beyond what Honeywell provides.',
    },
    {
        'id': 'DEV-006', 'title': 'Temperature Rating - Standard 60 deg C vs. Design 106 deg C',
        'severity': 'HIGH',
        'rfq_ref': 'S1.3 Process Data',
        'requirement': 'Design temperature: -8 deg C to 106 deg C.',
        'offered': 'Standard approved up to 60 deg C. High Temperature option listed but inclusion in price unclear.',
        'deviation': 'Standard prover rated to 60 deg C. Design temperature 106 deg C requires high-temp option.',
        'risk_owner': 'Honeywell (to confirm inclusion)',
        'resolution': 'Honeywell to confirm high-temperature option is included in quoted price. If adder, provide pricing. Operating temp 55-57 deg C is within range, but design temp 106 deg C is not.',
    },
    {
        'id': 'DEV-007', 'title': 'Modular Split - Module Dimensions Exceed RFQ Envelope',
        'severity': 'HIGH',
        'rfq_ref': 'S1.7 Envelope',
        'requirement': 'Max module: W 1.4m x L 2.56m x H 2.2m',
        'offered': 'Two-piece frame. ~2.7m per half (exceeds L), 1.6m width (exceeds W). Flow tube 2.6-2.7m, 3500 kg.',
        'deviation': 'Frame halves and flow tube exceed stated maximum dimensions.',
        'risk_owner': 'Guidant/Equinor (must accept), Sifab (execution), Honeywell (design)',
        'resolution': 'Guidant verbally accepted (18 Mar). Formal deviation request needed with detailed split-module drawings.',
    },
    {
        'id': 'DEV-008', 'title': 'SAT / Commissioning / Third-Party Inspection - Not Included',
        'severity': 'HIGH',
        'rfq_ref': 'S1.12 Testing Prover',
        'requirement': 'Gravimetric calibration witnessed. Water draw at factory and after re-assembly.',
        'offered': 'SAT, commissioning and third-party inspection not included. FAT at Tempe, AZ included.',
        'deviation': 'No SAT, no commissioning, no third-party inspection, no re-assembly water draw on Snorre A.',
        'risk_owner': 'Must be scoped and priced',
        'resolution': 'Honeywell supervision for re-assembly and SAT to be quoted separately. Sifab must not accept responsibility for test results - HW supervisor must be present.',
    },
    {
        'id': 'DEV-009', 'title': 'Materials - TR2000 BD20X Compliance Unconfirmed',
        'severity': 'HIGH',
        'rfq_ref': 'S1.8 Material',
        'requirement': 'Wetted/pressure parts per Equinor TR2000 PCS BD20X. 6Mo tubing/fittings.',
        'offered': 'AISI 316/316L. No TR2000 reference. No 6Mo mentioned.',
        'deviation': 'Honeywell does not reference TR2000 BD20X. 6Mo tubing not addressed.',
        'risk_owner': 'Honeywell (material compliance)',
        'resolution': 'Honeywell to confirm 316/316L meets TR2000 BD20X. 6Mo tubing: Sifab free-issue (Aker BP approach).',
    },
    {
        'id': 'DEV-010', 'title': 'PMI Scope - Less Hardware Ambiguity',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.8 Material (TR1427)',
        'requirement': 'PMI per TR1427: 10% SS316; 100% duplex/6Mo.',
        'offered': 'PMI on pressurized wetted parts (less hardware).',
        'deviation': '"Less hardware" is ambiguous.',
        'risk_owner': 'Honeywell',
        'resolution': 'Honeywell to clarify PMI scope and confirm TR1427 compliance.',
    },
    {
        'id': 'DEV-011', 'title': 'Electrical - NORSOK E-001 / TR3023 / Cable Type',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.10 Electrical Requirements',
        'requirement': 'Cable glands per NORSOK E-001 and TR3023. BFOU type, halogen-free cables.',
        'offered': 'Not addressed in Honeywell proposal.',
        'deviation': 'No mention of NORSOK E-001, TR3023, BFOU cables. Aker BP: cabling had to be re-done.',
        'risk_owner': 'Sifab (if free-issue), Honeywell (if in their scope)',
        'resolution': 'Sifab to supply NORSOK-compliant cables, glands, cable trays as free-issue.',
    },
    {
        'id': 'DEV-012', 'title': 'Document Deliverables - Standard Package vs. RFQ',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.16 Documents',
        'requirement': '30+ documents with specific timelines (2/4/8 WAO, WD).',
        'offered': 'Standard package within 2 weeks of shipment. Additional docs at extra charge.',
        'deviation': 'Many RFQ documents likely outside Honeywell standard package.',
        'risk_owner': 'Honeywell (standard docs), Sifab (Norsok-specific)',
        'resolution': 'Map Honeywell standard package vs RFQ. Request pricing for gaps.',
    },
    {
        'id': 'DEV-013', 'title': 'Fabrication Procedures - Guidant Approval',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.11 Fabrication',
        'requirement': 'Procedures approved by Guidant prior to fabrication.',
        'offered': 'Not addressed.',
        'deviation': 'Honeywell does not commit to Guidant approval before fabrication.',
        'risk_owner': 'Schedule risk',
        'resolution': 'Include in PO: all procedures to be submitted for Guidant approval before manufacturing.',
    },
    {
        'id': 'DEV-014', 'title': 'Lifting Lugs - NORSOK R-002',
        'severity': 'MEDIUM',
        'rfq_ref': 'S1.13 Lifting',
        'requirement': 'Lifting lugs/points per NORSOK R-002.',
        'offered': 'Not addressed.',
        'deviation': 'To be clarified.',
        'risk_owner': 'To be clarified',
        'resolution': 'To be clarified during engineering phase.',
    },
    {
        'id': 'DEV-015', 'title': '100mm Insulation Clearance - Unconfirmed',
        'severity': 'LOW',
        'rfq_ref': 'S1.11 Fabrication',
        'requirement': 'SVP prepared for min. 100mm insulation. 100mm insulation around flow tube is required.',
        'offered': 'Not confirmed for 12" flanges and vent/drain nozzles.',
        'deviation': 'Flow tube must accommodate 100mm insulation. Clearance uncertain around 12" RTJ flanges and vent/drain nozzles. Frame dimensions must account for insulation thickness.',
        'risk_owner': 'Honeywell (design), Sifab (structural)',
        'resolution': 'Honeywell to confirm in GA drawings that all connections and flow tube allow 100mm insulation. Frame design must account for insulation thickness around flow tube.',
    },
    {
        'id': 'DEV-016', 'title': 'Hydrostatic Test Pressure - BD20X Reference',
        'severity': 'LOW',
        'rfq_ref': 'S1.12 Testing',
        'requirement': 'Hydrostatic: 1.5x max design pressure of PCS BD20X (= 73.5 barg).',
        'offered': 'Standard testing per API MPMS 4.2 and PED.',
        'deviation': 'Test pressure per BD20X not explicitly confirmed.',
        'risk_owner': 'Honeywell',
        'resolution': 'Honeywell to confirm hydrotest pressure meets 1.5x BD20X.',
    },
    {
        'id': 'DEV-017', 'title': 'Seraphin Can - Justervesenet Certification',
        'severity': 'LOW',
        'rfq_ref': 'S1 Scope',
        'requirement': 'Seraphin can certified by Justervesenet, in SS316 cabinet.',
        'offered': 'Sifab scope (not Honeywell).',
        'deviation': 'N/A (Sifab scope). Lead time must be managed.',
        'risk_owner': 'Sifab',
        'resolution': 'Sifab to source from Pemberton. Justervesenet certification. SS316 cabinet with forklift lifting.',
    },
    {
        'id': 'DEV-018', 'title': 'Payment Terms - 30% Down Payment',
        'severity': 'HIGH',
        'rfq_ref': 'Guidant T&C (PP-PS-13)',
        'requirement': 'Per Guidant General Terms & Conditions.',
        'offered': '30% non-refundable down payment. 30 days net.',
        'deviation': 'Honeywell requires 30% upfront. Guidant T&C may differ.',
        'risk_owner': 'Sifab (cash flow risk)',
        'resolution': 'Guidant milestone payments must cover Honeywell 30% down. Do not pre-finance. Back-to-back payment required.',
    },
]


def build_deviation_list():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc.add_heading('Deviation List', level=0)
    doc.add_heading('SP-01415 Snorre A - Small Volume Prover', level=1)

    add_header_table(doc, [
        ('RFQ No.', 'GM-8501-1447'),
        ('Sifab Project', 'SP-01415'),
        ('Client', 'Guidant (Measurement Solutions Norway AS)'),
        ('End Client', 'Equinor - Snorre A'),
        ('Honeywell Proposal', '10465986-O-1010834 R0, dated 20 March 2026'),
        ('Prepared by', 'Sifab AS'),
        ('Date', '2026-03-23'),
        ('Revision', '0'),
    ])

    doc.add_paragraph('')
    note = doc.add_paragraph()
    r = note.add_run('Note: ')
    r.bold = True
    note.add_run(
        'Sifab AS acts as system integrator and intermediary between Guidant and Honeywell. '
        'Deviations are documented to ensure full transparency. Sifab does not accept technical '
        'risk for items outside its scope - risk sits with the party responsible for the deliverable.'
    )

    for dev in DEVIATIONS:
        doc.add_page_break()
        doc.add_heading(f"{dev['id']}: {dev['title']}", level=2)

        sev_p = doc.add_paragraph()
        sev_run = sev_p.add_run(f"Severity: {dev['severity']}")
        sev_run.bold = True
        sev_run.font.color.rgb = severity_color(dev['severity'])

        t = doc.add_table(rows=7, cols=2)
        t.style = 'Light Grid Accent 1'
        fields = [
            ('RFQ Reference', dev['rfq_ref']),
            ('RFQ Requirement', dev['requirement']),
            ('Honeywell Offer', dev['offered']),
            ('Deviation', dev['deviation']),
            ('Risk Owner', dev['risk_owner']),
            ('Proposed Resolution', dev['resolution']),
            ('Status', 'Open'),
        ]
        for i, (k, v) in enumerate(fields):
            set_cell(t.rows[i].cells[0], k, bold=True)
            set_cell(t.rows[i].cells[1], v)

    # Summary
    doc.add_page_break()
    doc.add_heading('Summary', level=1)

    high = sum(1 for d in DEVIATIONS if d['severity'] == 'HIGH')
    med = sum(1 for d in DEVIATIONS if d['severity'] == 'MEDIUM')
    low = sum(1 for d in DEVIATIONS if d['severity'] == 'LOW')

    p = doc.add_paragraph()
    p.add_run(f'HIGH: {high}  |  MEDIUM: {med}  |  LOW: {low}  |  Total: {len(DEVIATIONS)}')

    st = doc.add_table(rows=len(DEVIATIONS) + 1, cols=4)
    st.style = 'Light Grid Accent 1'
    for i, h in enumerate(['#', 'Deviation', 'Severity', 'Status']):
        set_cell(st.rows[0].cells[i], h, bold=True)
    for i, dev in enumerate(DEVIATIONS):
        set_cell(st.rows[i + 1].cells[0], dev['id'])
        set_cell(st.rows[i + 1].cells[1], dev['title'])
        set_cell(st.rows[i + 1].cells[2], dev['severity'])
        set_cell(st.rows[i + 1].cells[3], 'Open')

    outpath = DEST / 'SP-01415_Deviation_List_Rev0.docx'
    doc.save(str(outpath))
    print(f'Saved: {outpath}')


def build_compliance_matrix():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)

    doc.add_heading('RFQ Compliance Matrix', level=0)
    doc.add_heading('SP-01415 Snorre A - Small Volume Prover', level=1)

    add_header_table(doc, [
        ('RFQ No.', 'GM-8501-1447 Rev. A'),
        ('Sifab Project', 'SP-01415'),
        ('Client', 'Guidant (Measurement Solutions Norway AS)'),
        ('End Client', 'Equinor - Snorre A'),
        ('Honeywell Proposal', '10465986-O-1010834 R0'),
        ('Date', '2026-03-23'),
        ('Revision', '0'),
    ])

    doc.add_paragraph('')
    legend = doc.add_paragraph()
    legend.add_run('Legend: ').bold = True
    legend.add_run('C = Comply | PC = Partially Comply | D = Deviate | CR = Clarification Required')
    legend2 = doc.add_paragraph()
    legend2.add_run('Responsible: ').bold = True
    legend2.add_run('HW = Honeywell | SF = Sifab AS | GU = Guidant/Equinor')

    # Compliance data: (section, items)
    sections = [
        ('S1 - Scope of Supply', [
            ('1.0-1', 'Supply 1x SVP incl. controller', 'C', 'HW', '', 'SVP085 model quoted'),
            ('1.0-2', 'Supply 1x certified Seraphin can', 'C', 'SF', '', 'Sourced from Pemberton'),
            ('1.0-3', 'Seraphin can cert. by Justervesenet', 'CR', 'SF', 'DEV-017', 'Lead time to confirm'),
            ('1.0-4', 'Seraphin can in SS316 cabinet', 'C', 'SF', '', 'Confirmed 21 Mar'),
            ('1.0-5', 'Water draw kit (solenoid + manual valves)', 'PC', 'HW/SF', '', 'HW quotes controller kit. Full scope TBC'),
            ('1.0-6', 'Split SVP per S1.7 envelope', 'D', 'HW/SF', 'DEV-007', 'Modules exceed envelope'),
            ('1.0-7', 'Disassembly/packing onshore Norway', 'C', 'HW/SF', '', 'HW supervision required'),
            ('1.0-8', 'Re-assembly on Snorre A', 'C', 'HW/SF', 'DEV-008', 'HW supervisor rates pending'),
            ('1.0-9', 'Warranty from site release', 'D', 'HW', 'DEV-005', 'HW: 24m from delivery'),
            ('1.0-10', 'Warranty min. 28 months', 'D', 'HW', 'DEV-005', 'HW standard is 24 months'),
            ('1.0-11', 'CE marking PED + ATEX', 'C', 'HW', '', 'Included'),
        ]),
        ('S1.1 - Standards & Requirements', [
            ('1.1-1', 'API MPMS Chapter 4.2', 'C', 'HW', '', 'Confirmed'),
            ('1.1-2', 'Maleforskriften', 'CR', 'HW/SF', '', 'Not referenced by HW'),
            ('1.1-3', 'PED 2014/68/EU', 'C', 'HW', '', 'Included'),
            ('1.1-4', 'ATEX 2014/34/EU', 'C', 'HW', '', 'Included'),
            ('1.1-5', 'Machinery Directive', 'CR', 'HW', '', 'Not explicitly referenced'),
            ('1.1-6', 'EMC Directive', 'CR', 'HW', '', 'Not explicitly referenced'),
            ('1.1-7', 'TR2000 BD20X', 'D', 'HW', 'DEV-009', 'Not referenced by HW'),
            ('1.1-8', 'TR0042 Coating', 'PC', 'SF', 'DEV-004', 'Sifab scope'),
            ('1.1-9', 'TR1427 PMI', 'PC', 'HW', 'DEV-010', 'Scope ambiguous'),
            ('1.1-10', 'TR1826 Welding', 'D', 'HW', 'DEV-003', 'ASME IX deviation'),
            ('1.1-11', 'TR3023 E&I offshore', 'D', 'SF', 'DEV-011', 'Sifab free-issue'),
            ('1.1-12', 'TR3032 Field instr.', 'CR', 'HW/SF', '', 'Not addressed'),
            ('1.1-13', 'NORSOK E-001', 'D', 'SF', 'DEV-011', 'Not addressed by HW'),
            ('1.1-14', 'NORSOK L-004', 'D', 'HW', 'DEV-003', 'See welding'),
            ('1.1-15', 'NORSOK M-101', 'CR', 'SF', '', 'Sifab structural scope'),
            ('1.1-16', 'NORSOK M-501', 'PC', 'SF', 'DEV-004', 'Sifab scope'),
            ('1.1-17', 'NORSOK M-601', 'D', 'HW', 'DEV-003', 'See welding'),
            ('1.1-18', 'NORSOK M-630 MDS', 'CR', 'HW', 'DEV-009', 'MDS reference needed'),
            ('1.1-19', 'NORSOK R-002 Lifting', 'D', 'SF', 'DEV-014', 'Sifab engineering'),
        ]),
        ('S1.2 - Area Classification', [
            ('1.2-1', 'ATEX Zone 1, IIA T3 min', 'C', 'HW', '', 'Confirmed'),
            ('1.2-2', 'IP66 minimum', 'D', 'HW', 'DEV-001', 'IP65/IP56 offered'),
        ]),
        ('S1.3 - Process & Design Data', [
            ('1.3-1', 'Flow rate 67-750 m3/h', 'C', 'HW', '', 'SVP085 covers range'),
            ('1.3-2', 'Min stroke time 1 sec', 'CR', 'HW', '', 'Not explicitly confirmed'),
            ('1.3-3', 'Design pressure 49 barg', 'C', 'HW', '', 'CL600 exceeds requirement'),
            ('1.3-4', 'Design temp -8 to 106 deg C', 'D', 'HW', 'DEV-006', 'Standard 60 deg C only'),
            ('1.3-5', 'Repeatability <= 0.020%', 'C', 'HW', '', 'Confirmed'),
            ('1.3-6', 'Noise/sizing/pressure drop report', 'CR', 'HW', '', 'Not mentioned'),
        ]),
        ('S1.5 - Motor', [
            ('1.5-1', '230 VAC, 3-phase, 60 Hz', 'PC', 'HW', 'DEV-002', '60 Hz verbally confirmed'),
            ('1.5-2', 'ATEX Ex de or Ex e', 'C', 'HW', '', 'Confirmed'),
        ]),
        ('S1.6 - Connections', [
            ('1.6-1', '12" ANSI B16.5 CL600 RTJ', 'C', 'HW', '', 'CL600 RTJ confirmed'),
            ('1.6-2', '1" drain/vent CL600 RTJ', 'CR', 'HW', '', 'To be confirmed on GA'),
            ('1.6-3', 'Orientation upward/upward', 'CR', 'HW', '', 'To be confirmed on GA'),
            ('1.6-4', 'Metric tubing, Hoke Gyrolok', 'CR', 'HW/SF', '', 'May be Sifab free-issue'),
            ('1.6-5', '3x process thermowells TR2000', 'CR', 'SF', '', 'Supply responsibility TBC'),
            ('1.6-6', '1x rod thermowell', 'CR', 'HW/SF', '', 'Scope assignment TBC'),
            ('1.6-7', '1x DB&B pressure take-off TR2000', 'CR', 'SF', '', 'B&B valve is Sifab-supplied'),
            ('1.6-8', 'No pressure transmitters', 'C', 'HW', '', 'Confirmed'),
            ('1.6-9', 'Guidant free-issue temp elements - vendor install', 'CR', 'HW/SF', '', 'Installation responsibility TBC'),
        ]),
        ('S1.7 - Envelope', [
            ('1.7-1', 'Max W1.4 x L2.56 x H2.2m', 'D', 'HW/SF', 'DEV-007', 'Modules exceed limits'),
            ('1.7-2', 'Lifting points per module', 'D', 'SF', 'DEV-014', 'Not yet developed'),
            ('1.7-3', 'Drawing size/weight biggest module', 'PC', 'HW', '', 'Frame drawing provided, split details pending'),
        ]),
        ('S1.8 - Material', [
            ('1.8-1', 'Wetted parts per TR2000 BD20X', 'D', 'HW', 'DEV-009', '316/316L - TR2000 unconfirmed'),
            ('1.8-2', 'Flow tube min SS316 per M-630 MDS', 'C', 'HW', '', '316 SST confirmed'),
            ('1.8-3', 'Sour service NACE MR0175', 'C', 'HW', '', 'NACE stated'),
            ('1.8-4', 'Atmospheric parts min SS316', 'C', 'HW', '', '316 SST Offshore'),
            ('1.8-5', 'Structural frame SS316L', 'C', 'HW', '', 'Confirmed by Sidney'),
            ('1.8-6', 'Controller/interface box SS316', 'CR', 'HW', '', 'Not explicitly stated'),
            ('1.8-7', 'Dissimilar metals PTFE separation', 'CR', 'HW', '', 'Not addressed'),
            ('1.8-8', 'Cable glands SS316/NiPl brass', 'CR', 'HW/SF', 'DEV-011', 'Not addressed'),
            ('1.8-9', '6Mo tubing/fittings per TR2000', 'D', 'SF', 'DEV-009', 'Sifab free-issue'),
            ('1.8-10', 'Instrument valves SS316', 'CR', 'SF', '', 'Sifab scope (B&B valve)'),
            ('1.8-11', 'PMI per TR1427', 'PC', 'HW', 'DEV-010', 'Scope ambiguous'),
        ]),
        ('S1.9 - Painting', [
            ('1.9-1', 'TR0042 and NORSOK M-501', 'PC', 'SF', 'DEV-004', 'Sifab scope'),
            ('1.9-2', 'Flow tube System 6C', 'PC', 'SF', 'DEV-004', 'Painting in Norway'),
            ('1.9-3', 'SS structure: N/A', 'C', '', '', 'No painting needed'),
            ('1.9-4', 'Parts under insulation painted', 'C', 'SF', '', 'Confirmed 23 Mar'),
        ]),
        ('S1.10 - Electrical', [
            ('1.10-1', 'Cable glands NORSOK E-001/TR3023', 'D', 'SF', 'DEV-011', 'Sifab free-issue'),
            ('1.10-2', 'BFOU cables, halogen-free', 'D', 'SF', 'DEV-011', 'Sifab free-issue'),
        ]),
        ('S1.11 - Fabrication', [
            ('1.11-1', 'Welding per TR1826 / NORSOK M-601', 'D', 'HW', 'DEV-003', 'ASME IX deviation'),
            ('1.11-2', 'Piping per NORSOK L-004', 'D', 'HW', 'DEV-003', 'See above'),
            ('1.11-3', 'Structural per NORSOK M-101', 'CR', 'SF', '', 'Sifab structural scope'),
            ('1.11-4', 'Procedures approved by Guidant', 'D', 'HW', 'DEV-013', 'Not addressed by HW'),
            ('1.11-5', 'Prepared for 100mm insulation', 'CR', 'HW', 'DEV-015', '100mm insulation around flow tube required. Clearance for flanges unconfirmed.'),
            ('1.11-6', 'E&I per TR3023/I-001/E-001', 'D', 'SF', 'DEV-011', 'Sifab scope'),
        ]),
        ('S1.12 - Testing', [
            ('1.12-1', 'Gravimetric calibration at factory', 'C', 'HW', '', 'FAT at Tempe, AZ'),
            ('1.12-2', 'Witnessed by Buyer/Contractor/End Client/Authorities', 'PC', 'HW', 'DEV-008', 'FAT witness possible. Authorities TBC'),
            ('1.12-3', 'Water draw at factory with Seraphin can', 'C', 'HW/SF', '', 'Can sent to TruStop (confirmed 23 Mar)'),
            ('1.12-4', 'Repeatability <= 0.020%', 'C', 'HW', '', 'Confirmed'),
            ('1.12-5', 'Water draw after re-assembly Snorre A', 'D', 'HW/SF', 'DEV-008', 'SAT not included'),
            ('1.12-6', 'Hydrostatic 1.5x BD20X design pressure', 'CR', 'HW', 'DEV-016', 'Test pressure TBC'),
        ]),
        ('S1.13-1.15 - Lifting, Tagging, Spares', [
            ('1.13-1', 'Lifting per NORSOK R-002', 'CR', 'SF', 'DEV-014', 'To be clarified'),
            ('1.14-1', 'SS316 nameplate with CE + tag', 'CR', 'HW', '', 'Not explicitly addressed'),
            ('1.14-2', 'Individual SS316 tag signs', 'CR', 'HW/SF', '', 'TBC'),
            ('1.15-1', 'Commissioning + 2-year spares (option)', 'C', 'HW', '', '2-year kit included'),
        ]),
    ]

    for section_title, items in sections:
        doc.add_page_break()
        doc.add_heading(section_title, level=2)

        t = doc.add_table(rows=len(items) + 1, cols=6)
        t.style = 'Light Grid Accent 1'
        t.autofit = True

        hdrs = ['Ref', 'Requirement', 'Status', 'Resp.', 'DEV', 'Remarks']
        for i, h in enumerate(hdrs):
            set_cell(t.rows[0].cells[i], h, bold=True, size=8)

        for row_idx, item in enumerate(items):
            for col_idx, val in enumerate(item):
                set_cell(t.rows[row_idx + 1].cells[col_idx], val, size=8)

    # Summary
    doc.add_page_break()
    doc.add_heading('Compliance Summary', level=1)

    counts = {'C': 0, 'PC': 0, 'D': 0, 'CR': 0}
    for _, items in sections:
        for item in items:
            status = item[2]
            counts[status] = counts.get(status, 0) + 1

    total = sum(counts.values())
    summary = doc.add_table(rows=5, cols=3)
    summary.style = 'Light Grid Accent 1'
    set_cell(summary.rows[0].cells[0], 'Status', bold=True)
    set_cell(summary.rows[0].cells[1], 'Count', bold=True)
    set_cell(summary.rows[0].cells[2], '%', bold=True)
    for i, (s, label) in enumerate([('C', 'Comply'), ('PC', 'Partially Comply'), ('D', 'Deviate'), ('CR', 'Clarification Required')]):
        set_cell(summary.rows[i + 1].cells[0], f'{label} ({s})')
        set_cell(summary.rows[i + 1].cells[1], str(counts[s]))
        set_cell(summary.rows[i + 1].cells[2], f'{counts[s] * 100 // total}%')

    doc.add_paragraph('')
    doc.add_heading('Key Actions Required', level=2)
    actions = [
        'Honeywell must respond to 8 HIGH deviations (IP66, warranty, temperature, welding, materials, SAT, modules, payment)',
        '37% of requirements need clarification - most relate to Honeywell standard package vs. RFQ specifics',
        'Sifab scope items (painting, electrical, lifting, Seraphin can) need formal scoping',
        'Back-to-back terms must be established before PO to Honeywell',
        'Formal deviation requests needed for welding (ASME vs NORSOK) and module envelope',
    ]
    for a in actions:
        doc.add_paragraph(a, style='List Bullet')

    outpath = DEST / 'SP-01415_RFQ_Compliance_Matrix_Rev0.docx'
    doc.save(str(outpath))
    print(f'Saved: {outpath}')


if __name__ == '__main__':
    build_deviation_list()
    build_compliance_matrix()
