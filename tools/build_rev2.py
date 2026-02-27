import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0, 51, 102)

# ============================================================
# HEADER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SVP085 MODULAR SPLIT PLAN\nSnorre A Platform Access')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sifab AS  |  SP-01415  |  Rev 2  |  27 February 2026  |  INTERNAL')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('_' * 70)

# ============================================================
# 1. THE PROBLEM
# ============================================================
doc.add_heading('1. The Problem', level=1)

p = doc.add_paragraph()
p.add_run('The SVP085 must be installed on Snorre A but does not fit through the platform access door as a complete unit. ').font.size = Pt(10)
p.add_run('The frame is 48 mm too wide.').bold = True

t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['', 'SVP085', 'Door']
data = [
    ['Length', '5,258 mm', 'Max ~2,450 mm (corridor)'],
    ['Width (frame)', '1,448 mm', '1,400 mm  ← TOO WIDE'],
    ['Width (feet)', '1,588 mm', '1,400 mm  ← TOO WIDE'],
    ['Height', '1,316 mm (1,367 offshore)', '2,200 mm  ← FITS'],
]
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        t.rows[i+1].cells[j].text = val

p = doc.add_paragraph()
p.add_run('\nSolution: ').bold = True
p.add_run('Split the frame into 3 sections. Flow tube shipped from TruStop in wooden crate, fitted with transport wheels on platform, and rolled through the door. Frame sections flipped 90\u00b0 on their side (756 mm wide) to pass through the door.')

# ============================================================
# 2. ACCESS CONFIRMED BY GUIDANT
# ============================================================
doc.add_heading('2. Platform Access Confirmed (Guidant, 27 Feb 2026)', level=1)

p = doc.add_paragraph()
p.add_run('Guidant (Torleif Espegard) provided site photos confirming:')
doc.add_paragraph('Welding is OK on platform (preferred before prover in position)', style='List Bullet')
doc.add_paragraph('1.59 m width can work if ladder east of metering package is temporarily removed', style='List Bullet')
doc.add_paragraph('Transport route confirmed: through door, along corridor, turn west to installation', style='List Bullet')
doc.add_paragraph('Customer asks: "What do you envision for length and height?"', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Our answer: ').bold = True
p.add_run('Max 2,450 mm length per frame module. Frame sections ~756 mm high when flipped on side. Flow tube in wooden crate on wheels (~600-800 mm high).')

doc.add_paragraph('')
doc.add_paragraph('Corridor photos from Snorre A:').bold = True

# Add photos - 2 per row using a table
photo_table = doc.add_table(rows=2, cols=2)
photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER

photos = [
    ('C:/Users/SondreFalch/SifabAS/downloads/image001.png', 'Existing metering skid & size template'),
    ('C:/Users/SondreFalch/SifabAS/downloads/image002.png', 'Access door (1.4m frame, 1.59m possible)'),
    ('C:/Users/SondreFalch/SifabAS/downloads/image003.png', 'Corridor length view'),
    ('C:/Users/SondreFalch/SifabAS/downloads/image004.png', 'Transport routing toward installation'),
]

for idx, (path, caption) in enumerate(photos):
    row = idx // 2
    col = idx % 2
    cell = photo_table.cell(row, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(2.8))
    p2 = cell.add_paragraph(caption)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(8)
    p2.runs[0].font.italic = True

# ============================================================
# 3. SPLIT CONCEPT
# ============================================================
doc.add_heading('3. Split Concept', level=1)

p = doc.add_paragraph()
p.add_run('The prover is split into transport modules. The frame is cut into 3 sections at the bolted splice joints shown below. Hydraulic drive and controller/piping are removed and transported separately. Two scenarios are proposed for the flow tube:')

doc.add_paragraph('')

# Add split illustration
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('C:/Users/SondreFalch/SifabAS/downloads/SVP085_split_illustration.png', width=Inches(6.2))

p = doc.add_paragraph('Annotated from Honeywell GA Drawing (Manual Part No. 44200001). Red = splice joints. Green = additional mounting brackets.')
p.runs[0].font.size = Pt(8)
p.runs[0].font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ---- SCENARIO A ----
doc.add_heading('Scenario A: Flow tube as-is in wooden crate (recommended)', level=2)

p = doc.add_paragraph()
p.add_run('Flow tube shipped complete from TruStop in standard wooden crate. On platform, transport wheels fitted under crate and rolled straight through the door. Simplest and lowest risk option.').font.size = Pt(10)

t = doc.add_table(rows=4, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Item', 'Dimensions', 'Notes']):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
sc_a = [
    ['Flow tube + end flanges in crate', '~2,900 x 600 x 600 mm', 'One crate, ~1,700-2,200 kg'],
    ['Piston/rod assembly', 'Packed inside same crate or separate', 'Removed from tube for protection'],
    ['Transport wheels', 'Fitted under crate on platform', 'Sifab supplies'],
]
for i, row_data in enumerate(sc_a):
    for j, val in enumerate(row_data):
        t.rows[i+1].cells[j].text = val

p = doc.add_paragraph()
p.add_run('Crate is ~2.9 m long (exceeds 2,450 mm frame module limit) but only ~600 mm wide \u2014 easily maneuverable through the 1.4 m door and along the corridor.').font.size = Pt(10)

doc.add_paragraph('')

# ---- SCENARIO B ----
doc.add_heading('Scenario B: Remove end flanges + piston, pack in crates', level=2)

p = doc.add_paragraph()
p.add_run('End flanges and piston removed from flow tube at Sifab workshop. Tube body (~2,600 mm), flanges, and piston packed into 2 wooden crates. Tube crate is ~300 mm shorter than Scenario A and lighter. Flanges and RTJ faces protected separately.').font.size = Pt(10)

t = doc.add_table(rows=4, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Item', 'Dimensions', 'Notes']):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
sc_b = [
    ['Crate 1: Tube body (no flanges)', '~2,700 x 500 x 500 mm', 'Lighter, narrower (~1,500-1,800 kg)'],
    ['Crate 2: End flanges + piston + seals', '~800 x 600 x 600 mm', 'RTJ faces protected separately (~200-400 kg)'],
    ['Transport wheels', 'Fitted under each crate on platform', 'Sifab supplies'],
]
for i, row_data in enumerate(sc_b):
    for j, val in enumerate(row_data):
        t.rows[i+1].cells[j].text = val

p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run('Removing end flanges requires Honeywell procedure. RTJ flange faces must be protected. Reassembly of flanges on platform adds offshore man-hours. Piston seal inspection required before reinstallation.')

doc.add_paragraph('')

# ---- COMMON MODULES (frame + drive + controller) ----
doc.add_heading('Common Modules (both scenarios)', level=2)

t = doc.add_table(rows=6, cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Module', 'Contents', 'Size (transport)', 'Weight']
data = [
    ['2A', 'Frame - non-drive end (flipped 90\u00b0)', '~1,950 x 800 x 1,448 mm', '~400-500 kg'],
    ['2B', 'Frame - center (flipped 90\u00b0)', '~1,750 x 800 x 1,448 mm', '~400-500 kg'],
    ['2C', 'Frame - drive end (flipped 90\u00b0)', '~1,558 x 800 x 1,448 mm', '~400-500 kg'],
    ['2D', 'Hydraulic drive unit', '~1,200 x 800 x 800 mm', '~500-800 kg'],
    ['3', 'Controller + piping', 'Various (< 1,200 mm wide)', '~400-700 kg'],
]
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        t.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# Key points
doc.add_heading('Key Engineering Points', level=2)
doc.add_paragraph('Flow tube body is ~2,600 mm (calibrated bore between detector switches: 2,954 mm). Transported as one piece in both scenarios \u2014 the tube is NOT cut or shortened.', style='List Bullet')
doc.add_paragraph('Frame sections flipped 90\u00b0 on their side for transport (cross-section 756 x 1,448 mm). All frame modules < 2,450 mm.', style='List Bullet')
doc.add_paragraph('Hydraulic drive removed from frame to keep sections light and balanced.', style='List Bullet')
doc.add_paragraph('Bolted splice joints with precision dowel pins or machined register faces for \u00b10.5 mm alignment.', style='List Bullet')
doc.add_paragraph('Additional mounting brackets: 1 at far left end (non-drive end) + 2 at splice points = 3 new + 2 existing = 5 total. The left end bracket must be heavy-duty \u2014 the motor drives the piston back and forth through the flow tube, creating significant axial and dynamic loads at this end. This bracket must absorb the reaction forces from piston cycling.', style='List Bullet')
doc.add_paragraph('Water draw test (SAT) mandatory after reassembly to verify \u22640.020% repeatability.', style='List Bullet')

# ============================================================
# 4. REFERENCE DRAWING
# ============================================================
doc.add_heading('4. Reference: SVP085 Dimensions (Honeywell Manual)', level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('C:/Users/SondreFalch/SifabAS/downloads/prover_manual_images/page32_full.png', width=Inches(5.5))

p = doc.add_paragraph('Table 21 from Honeywell Enraf Manual Part No. 44200001')
p.runs[0].font.size = Pt(8)
p.runs[0].font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 5. HONEYWELL SCOPE & TQs
# ============================================================
doc.add_heading('5. Honeywell Scope (to be included in bid)', level=1)

p = doc.add_paragraph()
p.add_run('The modular split is Honeywell\'s scope of delivery. ').bold = True
p.add_run('Sifab provides local workshop, offshore logistics, and Norsok oversight. Honeywell owns the engineering, disassembly, reassembly, SAT, and warranty.')

doc.add_paragraph('')
doc.add_heading('What Honeywell must include:', level=2)

items = [
    'Frame designed with bolted splice joints from factory (not cut after build)',
    'Additional flow tube saddle supports at each splice location',
    'Flow tube: Scenario A \u2014 ship complete in wooden crate from TruStop (valve isolation, flange protection, VCI). Scenario B \u2014 remove end flanges + piston, pack in 2 crates. Both: crate must be suitable for transport wheels on platform.',
    'Hydraulic quick-disconnect couplings for drive removal/reinstallation',
    'Laser-marked alignment reference points at factory',
    'Complete disassembly/reassembly procedure with torque values and tolerances',
    'Honeywell personnel for disassembly (Sifab workshop, Sandnes) and reassembly (Snorre A)',
    'Water draw SAT after reassembly, witnessed by all parties + Justervesenet',
    'Warranty starts after successful SAT on Snorre A (min 28 months from commissioning)',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

# ============================================================
# 6. TECHNICAL QUERIES
# ============================================================
doc.add_heading('6. Technical Queries for Honeywell', level=1)

tqs = [
    ('TQ-009', 'CRITICAL', 'Provide detailed GA drawing of SVP085 with frame cross-section and foot details. Can frame width be reduced by removing feet?'),
    ('TQ-010', 'CRITICAL', 'Has Honeywell ever split an SVP for restricted offshore access? Share procedure/lessons learned.'),
    ('TQ-011', 'HIGH', 'Exact dry weight of SVP085 CL600 RTJ, broken down by component (flow tube, frame, drive, controller).'),
    ('TQ-012', 'CRITICAL', 'Can the flow tube be removed and reinstalled without affecting calibration? What alignment tolerances?'),
    ('TQ-013', 'HIGH', 'Will Honeywell warrant the split prover? Conditions?'),
    ('TQ-016', 'CRITICAL', 'Can the frame be designed from the start with bolted splice joints (design-for-split)?'),
    ('TQ-020', 'CRITICAL', 'Can Honeywell add saddle-type mounting points at each splice (4 total) without affecting calibration?'),
    ('TQ-021', 'CRITICAL', 'Exact flow tube length (bore only)? Wooden crate dimensions for transport? Crate must be suitable for fitting transport wheels underneath.'),
]

t = doc.add_table(rows=len(tqs)+1, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['TQ#', 'Priority', 'Question']):
    cell = t.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for i, (tq, prio, question) in enumerate(tqs):
    t.rows[i+1].cells[0].text = tq
    t.rows[i+1].cells[1].text = prio
    t.rows[i+1].cells[2].text = question

# ============================================================
# 7. NEXT STEPS
# ============================================================
doc.add_heading('7. Next Steps', level=1)

steps = [
    ('Send this plan + TQs to Honeywell (via Sidney Swart)', 'Tom', 'Before 4 March bid'),
    ('Meeting with Honeywell to discuss split feasibility', 'Tom / Sondre', 'ASAP'),
    ('Reply to Guidant: module dimensions (2,450 mm L x 756 mm H flipped) + confirm ~3.1 m crate can maneuver through corridor', 'Tom', 'This week'),
    ('Honeywell to provide split concept drawing + weight breakdown + warranty terms', 'Honeywell', 'With bid'),
]

t = doc.add_table(rows=len(steps)+1, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Action', 'Owner', 'Deadline']):
    cell = t.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for i, (action, owner, deadline) in enumerate(steps):
    t.rows[i+1].cells[0].text = action
    t.rows[i+1].cells[1].text = owner
    t.rows[i+1].cells[2].text = deadline

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph('_' * 70)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sifab AS  |  Bedriftsveien 20, 4313 Sandnes  |  www.sifab.no')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save
output_path = 'C:/Users/SondreFalch/SifabAS/projects/snorre-a-compact-prover/engineering/Modular_Split_Execution_Plan_SVP085_Rev2.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print('Done!')
